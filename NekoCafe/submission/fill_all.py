#!/usr/bin/env python3
"""
Generic template filler for NekoCafe D-series documents.
Reads content from JSON data files and fills in Word templates.
"""

import json
import os
from docx import Document
from docx.shared import Pt
from copy import deepcopy

TEMPLATE_DIR = "/Users/taohaoran/Documents/git/Software-Engineering/产出模板"
OUTPUT_DIR = "/Users/taohaoran/Documents/git/Software-Engineering/NekoCafe/submission"

# ── Helper functions ──

def set_cell(cell, text):
    """Replace cell content with single-paragraph text."""
    # Clear all paragraphs
    for p in cell.paragraphs:
        p.clear()
    # Remove extra paragraphs beyond the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)

def set_cell_multipara(cell, lines):
    """Replace cell content with multiple paragraphs."""
    for p in cell.paragraphs:
        p.clear()
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)

def add_table_rows(table, rows_data):
    """Add rows to an existing table."""
    for row_data in rows_data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            set_cell(row.cells[i], text)

# ── Fill a single document ──

def fill_document(doc_id, content):
    """Fill a template document with provided content."""
    # Determine template path
    template_map = {
        "D-01": "D-01_项目开题报告_模板.docx",
        "D-02": "D-02_需求规格说明书_SRS_模板.docx",
        "D-03": "D-03_概要设计说明书_模板.docx",
        "D-04": "D-04_详细设计说明书_模板.docx",
        "D-05": "D-05_数据库设计说明书_模板.docx",
    }

    output_map = {
        "D-01": "G03_T-01_D-01_项目开题报告_v1.0.docx",
        "D-02": "G03_T-01_D-02_需求规格说明书_SRS_v1.0.docx",
        "D-03": "G03_T-01_D-03_概要设计说明书_v1.0.docx",
        "D-04": "G03_T-01_D-04_详细设计说明书_v1.0.docx",
        "D-05": "G03_T-01_D-05_数据库设计说明书_v1.0.docx",
    }

    src = os.path.join(TEMPLATE_DIR, template_map[doc_id])
    dst = os.path.join(OUTPUT_DIR, output_map[doc_id])

    doc = Document(src)

    # ── Process metadata ──
    meta = content.get("metadata", {})
    if meta:
        meta_table = doc.tables[0]
        for row_idx, col_idx, value in meta:
            set_cell(meta_table.cell(row_idx, col_idx), value)

    # ── Process table cell replacements ──
    cells = content.get("cells", {})
    for table_idx_str, cell_data in cells.items():
        table_idx = int(table_idx_str)
        for cell_key, cell_value in cell_data.items():
            parts = cell_key.split(",")
            row_idx = int(parts[0])
            col_idx = int(parts[1])
            is_multipara = len(parts) > 2 and parts[2] == "multi"
            if is_multipara:
                set_cell_multipara(doc.tables[table_idx].cell(row_idx, col_idx), cell_value)
            else:
                set_cell(doc.tables[table_idx].cell(row_idx, col_idx), cell_value)

    # ── Process table row additions ──
    add_rows = content.get("add_rows", {})
    for table_idx_str, rows_data in add_rows.items():
        table_idx = int(table_idx_str)
        add_table_rows(doc.tables[table_idx], rows_data)

    doc.save(dst)
    print(f"{doc_id}: saved to {dst}")

# ── Main ──

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python3 fill_all.py <doc_id>")
        print("  doc_id: D-01, D-02, D-03, D-04, D-05, or ALL")
        sys.exit(1)

    target = sys.argv[1]

    # Load all content files
    if target == "ALL":
        targets = ["D-01", "D-02", "D-03", "D-04", "D-05"]
    else:
        targets = [target]

    for doc_id in targets:
        json_path = os.path.join(OUTPUT_DIR, f"content_{doc_id.lower()}.json")
        if not os.path.exists(json_path):
            print(f"Warning: {json_path} not found, skipping {doc_id}")
            continue
        with open(json_path, "r", encoding="utf-8") as f:
            content = json.load(f)
        fill_document(doc_id, content)

    print("Done!")
