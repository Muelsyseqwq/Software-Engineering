#!/usr/bin/env python3
"""Fill D-03 概要设计说明书 - preserves all formatting, uses 「」 for Chinese quotes."""

from docx import Document
from docx.shared import Pt
from copy import deepcopy
import glob, os

matches = glob.glob("/Users/taohaoran/Documents/git/Software-Engineering/产出模板/D-03_*模板.docx")
SRC = matches[0]
DST = "/Users/taohaoran/Documents/git/Software-Engineering/NekoCafe/submission/G03_T-01_D-03_概要设计说明书_v1.0.docx"

doc = Document(SRC)

def cs(cell, text):
    """Set cell text (single paragraph)."""
    for p in cell.paragraphs: p.clear()
    for p in cell.paragraphs[1:]: p._element.getparent().remove(p._element)
    cell.paragraphs[0].text = ""
    r = cell.paragraphs[0].add_run(text)
    r.font.name = "Microsoft YaHei"; r.font.size = Pt(10)

def cm(cell, lines):
    """Set cell text (multi-paragraph)."""
    for p in cell.paragraphs: p.clear()
    for p in cell.paragraphs[1:]: p._element.getparent().remove(p._element)
    for i, line in enumerate(lines):
        if i == 0: p = cell.paragraphs[0]
        else:
            new_el = deepcopy(cell.paragraphs[0]._element)
            cell._element.append(new_el)
            p = cell.paragraphs[i]
        p.clear()
        r = p.add_run(line)
        r.font.name = "Microsoft YaHei"; r.font.size = Pt(10)

# ── Cover metadata (T0) ──
t = doc.tables[0]
cs(t.cell(4,1), "第 3 组")
cs(t.cell(5,1), "计算机22-3 / 221234301 / 张某某\n计算机22-3 / 221234302 / 李某某\n计算机22-3 / 221234303 / 王某某\n计算机22-3 / 221234304 / 赵某某")
cs(t.cell(6,1), "2026 - 06 - 13")

# ═════════════ D-03 TABLE MAP ═════════════
# T0: metadata | T1: guidance | T2: cover info rows
# T3: section 1 heading area | T4: section 1 body (placeholder)
# T5: section 2.1 | T6: section 2.2 | T7: section 2.3 (guided) | T8: section 2 footer
# T9: section 3.1 (guided) | T10: section 3.2 | T11: section 3 footer
# T12: section 4.1 | T13: section 4.2 | T14: section 4 footer
# T15: section 5 (guided) | T16: section 6
# T17: section 7 guided | T18, T19: ADR areas | T20: section 7 footer
# T21: section 8.1 (guided) | T22: section 8.2 | T23: section 8 footer
# T24: section 9 (guided) | T25: section 10

# Section 1: 引言 (T4 = body)
cm(doc.tables[4].cell(0,0), [
    "1.1 编写目的与读者",
    "本文档是 NekoCafé 智慧猫咖餐饮预约平台的概要设计说明书，描述系统的总体架构、模块划分、技术选型、关键流程和部署方案。",
    "",
    "目标读者：",
    "• 开发团队：理解系统整体架构和模块职责，指导详细设计和编码",
    "• 评审老师：快速了解系统的技术方案和架构决策",
    "• 测试团队：理解系统结构以制定测试策略",
    "• 运维人员：了解部署架构和环境依赖",
    "",
    "1.2 与 SRS 的对应关系",
    "本文档的设计决策追溯到 D-02 SRS 的以下章节：",
    "• 总体架构 (§2.3 技术选型) → D-02 §2.4-2.6 运行环境与约束",
    "• 模块划分 (§3.1) → D-02 §4.1 功能列表（6 必做模块 + 选做）",
    "• 运行视图 (§4) → D-02 §3.2 用例详述（UC-01 至 UC-05）",
    "• 接口设计 (§8) → D-02 §6 外部接口需求",
    "• 安全设计 (§10) → D-02 §5 非功能性需求（安全性）"
])

# Section 2: 总体设计
# T5 = 2.1 设计目标 (might need to find right table)
# T6 = 2.2 架构图, T7 = 2.3 技术选型
cm(doc.tables[7].cell(0,0), [
    "2.1 设计目标与原则",
    "",
    "设计目标：",
    "• 性能：核心接口 P95 ≤ 350ms，支持 1000 并发预约",
    "• 可演进：模块化单体架构，内部按领域分包，便于后续拆分为微服务",
    "• 可观测：统一的 traceId 传递、结构化日志、Actuator 健康检查",
    "• 可测试：分层架构使各层可独立测试，Service 层为核心测试重点",
    "• 安全：无状态 JWT 鉴权、方法级 RBAC、BCrypt 密码哈希、Secret 环境变量注入",
    "",
    "设计原则：",
    "• 前后端分离：前端 SPA + 后端 RESTful API，通过 JSON 通信",
    "• 领域驱动分包：后端按业务领域（auth/user/store/order/payment/...）组织代码",
    "• 分层架构：Controller → Service → Mapper，各层职责清晰",
    "• 统一响应：所有 API 返回 ApiResult<T> 统一格式 {code, message, data, traceId, timestamp}",
    "• 数据库版本化：所有 DDL/DML 变更通过 Flyway 迁移脚本管理",
    "",
    "2.2 系统总体架构图",
    "系统采用前后端分离 + 模块化单体后端架构（C4 Container 级别描述）：",
    "",
    "┌─────────────────────────────────────────────────────────┐",
    "│                    用户浏览器 (Chrome/FF/Edge)            │",
    "│  ┌───────────────────────────────────────────────────┐  │",
    "│  │        Vue 3 SPA (Nginx 静态服务)                   │  │",
    "│  │  Element Plus UI  │  Pinia Store  │  ECharts       │  │",
    "│  │  Vue Router (RBAC) │  Axios (JWT Interceptor)     │  │",
    "│  └───────────────────────────────────────────────────┘  │",
    "└──────────────────────┬──────────────────────────────────┘",
    "                       │ HTTPS (REST JSON)",
    "┌──────────────────────▼──────────────────────────────────┐",
    "│              Nginx (反向代理 + 静态资源)                   │",
    "│         /api/* → backend:8080   /uploads/* → backend    │",
    "└──────────────────────┬──────────────────────────────────┘",
    "                       │",
    "┌──────────────────────▼──────────────────────────────────┐",
    "│         Spring Boot 3.3.5 (JDK 17) - Backend             │",
    "│  ┌──────────────────────────────────────────────────┐   │",
    "│  │  Security Layer: JWT Filter + @PreAuthorize RBAC  │   │",
    "│  ├──────────────────────────────────────────────────┤   │",
    "│  │  Controller Layer: 16 Controllers (按角色组织)     │   │",
    "│  ├──────────────────────────────────────────────────┤   │",
    "│  │  Service Layer: 20+ Services (业务逻辑)            │   │",
    "│  ├──────────────────────────────────────────────────┤   │",
    "│  │  Mapper Layer: MyBatis-Plus (数据访问)             │   │",
    "│  └──────────────────────────────────────────────────┘   │",
    "└──────┬────────────────────────────┬─────────────────────┘",
    "       │                            │",
    "┌──────▼──────┐              ┌──────▼──────┐",
    "│  MySQL 8    │              │  Redis 7    │",
    "│  (InnoDB)   │              │  (可选缓存)  │",
    "│  Flyway 迁移 │              │             │",
    "└─────────────┘              └─────────────┘",
    "",
    "2.3 技术选型与理由",
    "┌────────────┬──────────────────────┬─────────────────────────────────┐",
    "│ 层次        │ 选型                  │ 理由                             │",
    "├────────────┼──────────────────────┼─────────────────────────────────┤",
    "│ 前端框架    │ Vue 3 + Vite + TS     │ 组合式API, 生态完善, 构建快      │",
    "│ UI库       │ Element Plus 2.8      │ 企业级Vue3组件, 表格/表单丰富    │",
    "│ 图表        │ ECharts 5.5           │ 强大的数据可视化, 多图表类型      │",
    "│ 后端框架    │ Spring Boot 3.3.5     │ 企业级Java框架, 自动配置, 生态广  │",
    "│ 安全框架    │ Spring Security + JWT │ 无状态鉴权, 方法级RBAC           │",
    "│ ORM        │ MyBatis-Plus 3.5.7    │ 自动CRUD, 乐观锁, 逻辑删除, 分页  │",
    "│ 数据库      │ MySQL 8.0 + InnoDB    │ 成熟稳定, 事务, 行锁, utf8mb4    │",
    "│ 数据库迁移  │ Flyway                │ SQL版本控制, 可追溯, 可复现       │",
    "│ 部署        │ Docker + Nginx        │ 一键启动, 环境一致, 反向代理      │",
    "│ 文档        │ SpringDoc OpenAPI     │ 自动生成Swagger UI, 在线调试     │",
    "│ AI集成      │ OpenAI兼容API         │ 可选启用, 自动fallback规则推荐    │",
    "│ 测试        │ JUnit5 + JaCoCo + k6  │ 单元测试+覆盖率+性能压测          │",
    "└────────────┴──────────────────────┴─────────────────────────────────┘"
])

# T8: Section 2 footer
cs(doc.tables[8].cell(0,0),
    "以上总体设计描述了系统的架构目标、C4 容器级架构图和技术选型理由，为后续模块详细设计提供框架。")

# Section 3: 系统结构 (T9, T10, T11)
cm(doc.tables[9].cell(0,0), [
    "3.1 模块划分",
    "后端按领域驱动设计（DDD）分为 16 个包，每个包对应一个业务领域或技术基础设施：",
    "",
    "┌─────────────────────┬──────────────────────────────────────────────┐",
    "│ 包 (Package)         │ 职责                                           │",
    "├─────────────────────┼──────────────────────────────────────────────┤",
    "│ config              │ 全局配置：WebMvc、MyBatis-Plus、OpenAPI、Health│",
    "│ common.exception    │ 统一异常：BizException + GlobalExceptionHandler │",
    "│ common.result       │ 统一响应：ApiResult<T> 泛型返回封装              │",
    "│ common.pagination   │ 统一分页：PageRequest/PageResponse             │",
    "│ security            │ 认证鉴权：JWT生成/校验、SecurityConfig、过滤器  │",
    "│ auth                │ 注册登录：注册、登录（用户名/手机/邮箱）、个人信息│",
    "│ user                │ 用户管理：User/Role/UserRole/MemberAccount实体  │",
    "│ admin               │ 总部管理：用户管理、角色管理、店长分配           │",
    "│ store               │ 门店管理：门店CRUD、桌位管理、LBS附近门店       │",
    "│ menu                │ 菜单管理：菜品分类、菜品CRUD、门店菜单查询       │",
    "│ reservation         │ 预约管理：时段生成、预约创建/取消、乐观锁库存    │",
    "│ queue               │ 排队管理：取号、叫号、入座、队列重置             │",
    "│ order               │ 订单管理：订单创建/取消、状态流转、库存扣减      │",
    "│ payment             │ 支付管理：沙箱支付、幂等键防重                   │",
    "│ customer            │ 顾客门户：首页、积分、偏好、评价、退款申请       │",
    "│ staff               │ 店员操作：签到、接单、桌位管理、排队操作         │",
    "│ manager             │ 店长管理：经营指标、员工排班、退款审批、活动确认  │",
    "│ dashboard           │ 数据看板：坪效/翻台率/复购率计算与趋势分析       │",
    "│ recommend           │ 推荐引擎：规则推荐 + AI文案（可选）              │",
    "│ cat                 │ 猫咪管理：猫咪档案、健康记录、体重趋势           │",
    "│ activity            │ 活动管理：活动CRUD、门店分发、顾客领取           │",
    "│ table               │ 桌位状态：桌位公共查询接口                      │",
    "└─────────────────────┴──────────────────────────────────────────────┘",
    "",
    "3.2 模块间依赖",
    "依赖关系遵循自上而下的分层原则：",
    "• Controller 层依赖 Service 层，不直接依赖 Mapper",
    "• Service 层依赖 Mapper 层和其他 Service（通过 Spring DI）",
    "• Mapper 层依赖 Entity，不依赖上层",
    "• common 包被所有业务包依赖",
    "• security 包被 web 层依赖（通过 Filter 拦截）",
    "• 跨领域调用通过 Service 层注入实现，避免循环依赖"
])

cs(doc.tables[11].cell(0,0),
    "以上系统结构描述了后端 16 个模块的职责划分和依赖关系，模块间低耦合、高内聚，便于团队并行开发和后续维护。")

# Section 4: 运行视图 (T12, T13, T14)
cm(doc.tables[12].cell(0,0), [  # T12 = 4.1 sequence diagrams (might be wrong index)
    "4.1 关键流程顺序图",
    "",
    "【流程1：顾客预约→支付→履约 完整流程】",
    "",
    "  Customer        前端SPA        Backend         MySQL",
    "    │               │               │               │",
    "    │ 1.选择门店     │               │               │",
    "    │──────────────>│               │               │",
    "    │               │ 2.GET /reservation/slots     │",
    "    │               │──────────────>│               │",
    "    │               │               │ 3.SELECT slots│",
    "    │               │               │──────────────>│",
    "    │               │               │<──────────────│",
    "    │               │<──────────────│               │",
    "    │ 4.选择时段+提交│               │               │",
    "    │──────────────>│               │               │",
    "    │               │ 5.POST /reservation          │",
    "    │               │──────────────>│               │",
    "    │               │               │ 6.UPDATE slot │",
    "    │               │               │ (乐观锁扣减)   │",
    "    │               │               │──────────────>│",
    "    │               │               │ 7.INSERT resv │",
    "    │               │               │──────────────>│",
    "    │               │<──reservationNo│               │",
    "    │<──预约成功     │               │               │",
    "    │               │               │               │",
    "    │ 8.浏览菜单+下单│               │               │",
    "    │──────────────>│               │               │",
    "    │               │ 9.POST /order │               │",
    "    │               │──────────────>│               │",
    "    │               │               │ 10.UPDATE dish │",
    "    │               │               │ (库存条件扣减)  │",
    "    │               │               │ 11.INSERT order│",
    "    │               │<──orderNo     │               │",
    "    │               │               │               │",
    "    │ 12.发起支付    │               │               │",
    "    │──────────────>│               │               │",
    "    │               │ 13.POST /payment/sandbox     │",
    "    │               │──────────────>│               │",
    "    │               │               │ 14.幂等查重    │",
    "    │               │               │ 15.INSERT pay  │",
    "    │               │               │ 16.UPDATE order│",
    "    │               │               │ 17.发放积分    │",
    "    │               │<──payment OK  │               │",
    "    │<──支付成功     │               │               │",
    "",
    "【流程2：店员履约流程】",
    "",
    "  STAFF          前端SPA        Backend         MySQL",
    "    │               │               │               │",
    "    │ 1.查看今日预约 │               │               │",
    "    │──────────────>│               │               │",
    "    │               │ 2.GET /staff/reservations/today",
    "    │               │──────────────>│               │",
    "    │               │<──预约列表     │               │",
    "    │<──展示        │               │               │",
    "    │               │               │               │",
    "    │ 3.顾客到店签到 │               │               │",
    "    │──────────────>│               │               │",
    "    │               │ 4.POST /staff/reservations/{id}/check-in",
    "    │               │──────────────>│               │",
    "    │               │               │ 5.校验门店归属  │",
    "    │               │               │ 6.UPDATE resv  │",
    "    │               │<──签到成功     │ (checked_in)  │",
    "    │               │               │               │",
    "    │ 7.开始制作订单 │               │               │",
    "    │──────────────>│               │               │",
    "    │               │ 8.POST /staff/orders/{id}/start",
    "    │               │──────────────>│               │",
    "    │               │               │ 9.UPDATE order │",
    "    │               │<──制作中       │ (PREPARING)   │",
    "    │               │               │               │",
    "    │ 10.完成订单    │               │               │",
    "    │──────────────>│               │               │",
    "    │               │ 11.POST /staff/orders/{id}/complete",
    "    │               │──────────────>│               │",
    "    │               │               │ 12.UPDATE order│",
    "    │               │<──已完成       │ (COMPLETED)   │"
])

# T13: 4.2 状态机
cm(doc.tables[13].cell(0,0), [
    "4.2 状态机",
    "",
    "【订单状态机】",
    "订单（FoodOrder）在其生命周期中经历以下状态流转：",
    "",
    "                    ┌─────────┐",
    "                    │ CREATED │ (初始状态：订单已创建，待支付)",
    "                    └────┬────┘",
    "                         │ pay() [沙箱支付]",
    "                         ▼",
    "                    ┌─────────┐",
    "              ┌─────│  PAID   │ (已支付，等待店员接单)",
    "              │     └────┬────┘",
    "              │          │ startOrder() [店员接单]",
    "              │          ▼",
    "              │     ┌───────────┐",
    "              │     │ PREPARING │ (制作中)",
    "              │     └─────┬─────┘",
    "              │           │ completeOrder() [店员完成]",
    "              │           ▼",
    "              │     ┌───────────┐",
    "              │     │ COMPLETED │ (已完成，可评价/可退款)",
    "              │     └───────────┘",
    "              │",
    "              │  cancel() [顾客取消]",
    "              ▼",
    "         ┌───────────┐",
    "         │ CANCELLED │ (已取消)",
    "         └───────────┘",
    "",
    "状态转移规则表：",
    "┌──────────────┬─────────────────┬──────────────┬──────────┬──────────────────┐",
    "│ 当前状态      │ 操作              │ 目标状态      │ 角色      │ 约束              │",
    "├──────────────┼─────────────────┼──────────────┼──────────┼──────────────────┤",
    "│ CREATED      │ 沙箱支付          │ PAID         │ CUSTOMER │ 订单归属本人       │",
    "│ CREATED      │ 取消订单          │ CANCELLED    │ CUSTOMER │ 未支付才可取消     │",
    "│ PAID         │ 店员接单           │ PREPARING    │ STAFF    │ 所属门店店员       │",
    "│ PAID         │ 申请退款           │ REFUNDING    │ CUSTOMER │ 未申请过退款       │",
    "│ PREPARING    │ 店员完成           │ COMPLETED    │ STAFF    │ 已接单            │",
    "│ PREPARING    │ 申请退款           │ REFUNDING    │ CUSTOMER │ 未申请过退款       │",
    "│ REFUNDING    │ 店长同意退款       │ REFUNDED     │ STORE_MGR│ 审核通过           │",
    "│ REFUNDING    │ 店长拒绝退款       │ REFUND_REJECTED│STORE_MGR│ 填写拒绝原因       │",
    "│ COMPLETED    │ 顾客评价           │ COMPLETED    │ CUSTOMER │ 订单归属本人       │",
    "└──────────────┴─────────────────┴──────────────┴──────────┴──────────────────┘",
    "",
    "【退款状态机】",
    "  REQUESTED ──approve()──> APPROVED ──> REFUNDED",
    "      │                                          ",
    "      └──reject()───> REJECTED                  ",
    "",
    "【桌位状态机】",
    "  AVAILABLE ←→ OCCUPIED ←→ NEEDS_CLEANING → AVAILABLE"
])

cs(doc.tables[14].cell(0,0),
    "以上运行视图通过 2 张顺序图和 3 个状态机描述了系统的核心业务流程和实体生命周期，覆盖了 SRS 中的主要用例（UC-01 至 UC-05）。")

# Section 5: 部署视图 (T15)
cm(doc.tables[15].cell(0,0), [
    "系统部署拓扑图（Docker Compose 一键部署）：",
    "",
    "┌──────────────────────────────────────────────────────────────┐",
    "│                    宿主机 (评委电脑)                           │",
    "│                                                              │",
    "│  ┌───────────────────┐  ┌───────────────────┐               │",
    "│  │  nginx:1.27-alpine │  │  backend:8080     │               │",
    "│  │  (frontend:80)     │  │  Spring Boot Jar  │               │",
    "│  │  /        → SPA    │  │  /api/*           │               │",
    "│  │  /api/*   → proxy  │  │  /uploads/*       │               │",
    "│  │  /uploads/*→ proxy │  │                   │               │",
    "│  └─────────┬─────────┘  └────────┬──────────┘               │",
    "│            │                     │                           │",
    "│            │                     │                           │",
    "│  ┌─────────▼─────────────────────▼──────────┐               │",
    "│  │              mysql:8.0                    │               │",
    "│  │         (3306, InnoDB, utf8mb4)           │               │",
    "│  │         Flyway 自动迁移                    │               │",
    "│  └───────────────────────────────────────────┘               │",
    "│                                                              │",
    "│  ┌───────────────────┐                                      │",
    "│  │  redis:7-alpine   │  (可选，用于缓存/Session)             │",
    "│  └───────────────────┘                                      │",
    "└──────────────────────────────────────────────────────────────┘",
    "",
    "网络边界：",
    "• 对外暴露端口：80 (Nginx/前端)",
    "• 内部网络：backend、mysql、redis 通过 Docker 内部网络通信",
    "• 数据库端口 3306 不对外暴露",
    "• Backend 端口 8080 仅 Nginx 内部访问",
    "",
    "数据流向：",
    "• 用户请求 → Nginx:80 → (静态资源) 直接返回 SPA 文件",
    "• 用户请求 → Nginx:80 → (API) proxy_pass → backend:8080 → MySQL:3306",
    "• 上传文件 → Nginx:80 → proxy_pass → backend:8080 → 本地磁盘 /uploads/"
])

# Section 6: 数据视图 (T16)
cm(doc.tables[16].cell(0,0), [
    "高层 ER 图（核心实体关系，详见 D-05 数据库设计说明书）：",
    "",
    "  user ──< user_role >── role",
    "  user ── member_account",
    "  user ── user_preference",
    "  user ──< reservation",
    "  user ──< food_order",
    "  user ──< review",
    "  user ──< refund_request",
    "  user ──< points_transaction",
    "",
    "  store ──< dining_table",
    "  store ──< reservation_slot >── dining_table",
    "  store ──< dish_category >──< dish",
    "  store ──< cat",
    "  store ──< user_store_role >── user",
    "  store ──< staff_shift >── user",
    "  store ──< waiting_queue_counter >──< waiting_queue_ticket",
    "",
    "  food_order ──< food_order_item >── dish",
    "  food_order ── payment_record",
    "",
    "  reservation ── dining_table",
    "  reservation ── reservation_slot",
    "",
    "  promotion_activity ──< activity_store >── store",
    "  promotion_activity ──< reward_catalog",
    "",
    "  cat ──< cat_health_record",
    "",
    "主要数据流：",
    "• 用户注册 → user + user_role + member_account",
    "• 创建预约 → reservation + reservation_slot(库存扣减)",
    "• 创建订单 → food_order + food_order_item + dish(库存扣减)",
    "• 支付 → payment_record + food_order(状态更新) + points_transaction",
    "• 退款 → refund_request → food_order(状态更新)",
    "• 看板查询 → 聚合 payment_record + food_order + reservation + store"
])

# Section 7: ADR (T17, T18, T19, T20)
cm(doc.tables[17].cell(0,0), [
    "以下记录 3 份关键架构决策：",
    "",
    "═══════════════════════════════════════════",
    "ADR-001：选择模块化单体而非微服务架构",
    "═══════════════════════════════════════════",
    "Status：Accepted",
    "Context：课程项目团队 4 人，开发周期 2 周。需要完成 6 个必做模块 + 选做功能。如果采用微服务架构，需要额外的服务发现、配置中心、API 网关、分布式事务等基础设施，开发和调试成本显著增加。",
    "Decision：采用模块化单体架构（Modular Monolith）。所有代码在一个 Spring Boot 应用中，但内部严格按领域分包（16 个 package），Controller/Service/Mapper 分层清晰，领域间通过 Service 层接口通信。",
    "Consequences：",
    "  • 正面：降低分布式复杂度；简化调试、测试和部署；团队可并行开发不同领域包；Docker 一键启动。",
    "  • 负面：所有模块共享一个数据库连接池；无法独立扩缩容；任意模块故障可能影响整个应用。",
    "  • 缓解：领域包之间低耦合；使用 Flyway 管理数据库变更；通过 Docker Compose 提供一致的运行环境。",
    "Alternatives Considered：",
    "  • 微服务架构：被否决——团队规模和周期不足以支撑分布式系统的基础设施和运维。",
    "  • 纯单体（无分包）：被否决——不利于团队并行开发和代码维护。",
    "",
    "═══════════════════════════════════════════",
    "ADR-002：选择 MySQL 8 + InnoDB 作为数据库",
    "═══════════════════════════════════════════",
    "Status：Accepted",
    "Context：系统需要存储用户、门店、预约、订单、支付、猫咪等结构化数据，需要事务支持和并发控制。课程设计要求数据库可一键部署，评委无需手动配置。",
    "Decision：选择 MySQL 8.0 + InnoDB 存储引擎。使用 utf8mb4 字符集支持 emoji。通过 Flyway 管理所有 DDL/DML 变更，确保数据库 schema 的可追溯性和可复现性。",
    "Consequences：",
    "  • 正面：成熟稳定；事务（ACID）支持；行级锁适合高并发预约场景；Docker Hub 有官方镜像，一键部署。",
    "  • 负面：垂直扩展有上限；相比 PostgreSQL 在 JSON 查询和地理空间功能上较弱。",
    "  • 缓解：课程设计数据量可控；LBS 距离计算在应用层（Haversine 公式）而非数据库层。",
    "Alternatives Considered：",
    "  • PostgreSQL：被否决——虽然功能更强，但团队更熟悉 MySQL，学习成本更低。",
    "  • H2 嵌入式数据库：被否决——不适合演示和生产模拟场景。",
    "",
    "═══════════════════════════════════════════",
    "ADR-003：选择 JWT 无状态鉴权而非 Session 鉴权",
    "═══════════════════════════════════════════",
    "Status：Accepted",
    "Context：系统有 5 种用户角色，需要前端 SPA + 后端 REST API 的鉴权方案。需要支持跨域（开发阶段 Vite dev server 端口 5173，后端端口 8080），且不希望后端维护 Session 状态。",
    "Decision：采用 JWT（JSON Web Token）无状态鉴权。用户登录后后端签发 JWT（含 userId、username、roles），前端存储在 localStorage，每次请求通过 Authorization: Bearer <token> 头传递。后端通过 OncePerRequestFilter 验证 token 有效性，解析出 AuthPrincipal 注入 Controller。角色鉴权通过 @PreAuthorize(hasRole(...)) 实现。",
    "Consequences：",
    "  • 正面：无状态，后端不需要 Session 存储；天然支持跨域和集群扩展；JJWT 库成熟。",
    "  • 负面：Token 一旦签发无法主动撤销（只能等过期）；Token 体积随 roles 增多而增大。",
    "  • 缓解：设置合理的过期时间（默认 120 分钟）；敏感操作二次校验；前端登出时清除 localStorage。",
    "Alternatives Considered：",
    "  • Session + Cookie：被否决——需要后端维护 Session，不支持跨域 SPA 场景，与无状态 REST 理念冲突。",
    "  • OAuth2/OIDC：被否决——课程项目无第三方登录需求，引入 OAuth2 基础设施过重。"
])

cs(doc.tables[20].cell(0,0),
    "以上 3 份 ADR 记录了系统在架构风格、数据库选型和鉴权方案三个关键维度上的决策过程与权衡，确保架构决策有据可查、可追溯。")

# Section 8: 接口设计 (T21, T22, T23)
cm(doc.tables[21].cell(0,0), [
    "8.1 对外接口（OpenAPI 摘要）",
    "系统通过 SpringDoc OpenAPI 自动生成 Swagger UI 文档（/swagger-ui.html）。以下是关键 API 端点摘要：",
    "",
    "【公开接口（无需鉴权）】",
    "• GET /api/health — 健康检查",
    "• POST /api/auth/login — 用户登录",
    "• POST /api/auth/register — 用户注册",
    "• GET /api/store — 门店列表",
    "• GET /api/store/{id} — 门店详情",
    "• GET /api/store/nearby — LBS 附近门店",
    "• GET /api/menu/stores/{storeId} — 门店菜单",
    "• GET /api/reservation/slots — 预约时段查询",
    "• GET /api/recommend/status — 推荐服务状态",
    "",
    "【CUSTOMER 角色接口】",
    "• POST /api/reservation — 创建预约",
    "• GET /api/reservation/me — 我的预约",
    "• POST /api/reservation/{id}/cancel — 取消预约",
    "• POST /api/order — 创建订单",
    "• GET /api/order/me — 我的订单",
    "• POST /api/order/{id}/cancel — 取消订单",
    "• POST /api/payment/sandbox — 沙箱支付",
    "• GET /api/customer/home — 顾客首页",
    "• GET /api/customer/activities — 活动列表",
    "• GET /api/customer/points — 我的积分",
    "• GET /api/customer/rewards — 奖励列表",
    "• POST /api/customer/rewards/{id}/redeem — 兑换奖励",
    "• GET/PUT /api/customer/preferences — 用户偏好",
    "• POST /api/customer/orders/{id}/reviews — 订单评价",
    "• POST /api/customer/orders/{id}/refunds — 申请退款",
    "• GET /api/customer/refunds/me — 我的退款",
    "• GET /api/recommend/customer — 门店推荐",
    "• GET /api/recommend/customer/cats — 猫咪推荐",
    "",
    "【STAFF 角色接口】",
    "• GET /api/staff/reservations/today — 今日预约",
    "• POST /api/staff/reservations/{id}/check-in — 预约签到",
    "• GET /api/staff/orders/pending — 待处理订单",
    "• POST /api/staff/orders/{id}/start — 开始制作",
    "• POST /api/staff/orders/{id}/complete — 完成订单",
    "• GET /api/staff/tables — 桌位列表",
    "• PUT /api/staff/tables/{id}/status — 更新桌位状态",
    "• POST /api/staff/queues/{storeId}/next — 叫号",
    "• POST /api/staff/queues/tickets/{ticketId}/seat — 入座",
    "",
    "【STORE_MANAGER 角色接口】",
    "• GET/PUT /api/manager/store — 门店信息与更新",
    "• GET /api/manager/metrics — 门店经营指标",
    "• GET /api/manager/orders — 订单管理",
    "• PUT /api/manager/orders/{id}/refund — 退款审批",
    "• GET /api/manager/cats — 猫咪管理",
    "• GET/POST/PUT /api/manager/staff — 员工管理",
    "• PUT /api/manager/staff/{id}/dismiss — 辞退员工",
    "• PUT /api/manager/dishes/{id}/price — 菜品定价",
    "",
    "【HQ_OPERATOR 角色接口】",
    "• GET /api/dashboard/summary — 看板摘要",
    "• GET /api/dashboard/revenue — 营收趋势",
    "• GET /api/dashboard/store-summary — 门店汇总",
    "• GET /api/dashboard/cross-store — 跨门店对比",
    "• GET/POST/PUT/DELETE /api/activity — 活动管理",
    "• POST /api/activity/{id}/publish — 发布活动",
    "• GET/POST/PUT/DELETE /api/admin/stores — 门店管理",
    "• GET/POST/DELETE /api/admin/store-managers — 店长分配",
    "",
    "【CAT_CARETAKER 角色接口】",
    "• GET/POST /api/cats — 猫咪列表与新增",
    "• GET/PUT/DELETE /api/cats/{id} — 猫咪详情/更新/删除",
    "• GET /api/cats/{id}/health-records — 健康记录",
    "• POST /api/cats/{id}/health-records — 添加健康记录",
    "• GET /api/cats/{id}/weight-trend — 体重趋势",
    "• PATCH /api/cats/{id}/health-status — 更新健康状态",
    "",
    "统一错误码字典：",
    "• 0: 成功",
    "• 1001: 用户名已存在",
    "• 1002: 账号或密码错误",
    "• 2001: 未登录 / Token 无效",
    "• 2002: 无访问权限",
    "• 3003: 预约人数超过容量",
    "• 3004: 时段已约满",
    "• 3005: 库存不足",
    "• 3006: 菜品已下架",
    "• 3007: 优惠券无效",
    "• 3008: 订单已支付",
    "• 5000: 系统内部错误"
])

# T22: 8.2 内部接口
cm(doc.tables[22].cell(0,0), [
    "8.2 内部接口（模块间）",
    "后端模块间通过 Spring 依赖注入进行同步调用，不使用消息队列或 RPC：",
    "",
    "关键内部调用关系：",
    "• AuthService → UserMapper, RoleMapper, UserRoleMapper, MemberAccountMapper",
    "• ReservationService → ReservationSlotMapper（乐观锁条件更新）",
    "• OrderService → DishMapper（库存条件扣减）, FoodOrderMapper, FoodOrderItemMapper",
    "• PaymentService → PaymentRecordMapper（幂等键查重）, FoodOrderMapper（状态更新）, MemberAccountMapper（积分发放）",
    "• StaffService → ReservationMapper, FoodOrderMapper, DiningTableMapper（校验门店归属）",
    "• StoreManagerService → 多个 Mapper（跨领域聚合查询）",
    "• RecommendService → AiReasonProperties（配置）, OpenAiCompatibleReasonClient（外部API）",
    "• CustomerService → UserPreferenceMapper, ReviewMapper, RefundRequestMapper, PointsTransactionMapper",
    "",
    "共享数据契约：",
    "• ApiResult<T>：所有 Controller 返回的统一响应格式",
    "• PageRequest/PageResponse：统一分页请求/响应",
    "• AuthPrincipal：JWT 解析后的用户身份（通过 @AuthenticationPrincipal 注入）",
    "• BizException：统一业务异常（code + message），由 GlobalExceptionHandler 处理"
])

cs(doc.tables[23].cell(0,0),
    "以上接口设计定义了系统全部 RESTful API（按角色分组）、统一错误码字典和模块间内部调用关系，为详细设计和前端对接提供完整接口契约。")

# Section 9: 错误处理与日志策略 (T24)
cm(doc.tables[24].cell(0,0), [
    "9.1 错误分类与处理",
    "",
    "错误分类：",
    "• 业务异常（BizException）：由业务逻辑抛出的可预期异常，携带业务错误码（1000-4999），GlobalExceptionHandler 捕获后返回 HTTP 400 + ApiResult（code=业务码, message=中文描述）",
    "• 校验异常（MethodArgumentNotValidException）：由 @Valid 触发的参数校验失败，GlobalExceptionHandler 捕获后返回 HTTP 400 + 具体字段错误",
    "• 鉴权异常（AccessDeniedException / AuthenticationException）：Spring Security 抛出的权限异常，返回 HTTP 401/403 + 中文提示",
    "• 系统异常（Exception）：未预期的运行时异常，GlobalExceptionHandler 兜底捕获，返回 HTTP 500 + ApiResult（code=5000, message=「系统繁忙，请稍后再试」）",
    "",
    "9.2 日志策略",
    "• 日志级别：",
    "  - ERROR：系统异常、数据库连接失败、外部 API 调用失败",
    "  - WARN：业务异常（如库存不足、权限不足）",
    "  - INFO：关键业务节点（注册、登录、创建预约、支付、退款审批）",
    "  - DEBUG：详细的 SQL 参数（开发环境开启，生产关闭）",
    "• traceId 传递：每个 HTTP 请求生成唯一 traceId（MDC 或 UUID），在 ApiResult.traceId 中返回，日志中记录，便于问题追踪",
    "• 关键告警阈值：",
    "  - 预约并发冲突率 > 10% → 需要增加时段库存或扩容",
    "  - 支付失败率 > 5% → 检查支付沙箱逻辑",
    "  - 数据库连接池等待时间 > 1s → 检查慢查询",
    "• 敏感信息脱敏：手机号中间 4 位、身份证号中间 8 位在日志中用 *** 替换"
])

# Section 10: 安全设计 (T25)
cm(doc.tables[25].cell(0,0), [
    "10 安全设计",
    "",
    "鉴权（Authentication）：",
    "• 所有非公开接口通过 JwtAuthenticationFilter 校验 Bearer Token",
    "• Token 使用 HMAC-SHA256 签名，Secret 通过环境变量 JWT_SECRET 注入",
    "• Token 有效期默认 120 分钟，过期后前端自动跳转登录页",
    "",
    "授权（Authorization）：",
    "• 方法级 RBAC：Controller 方法使用 @PreAuthorize(\\\"hasRole('ROLE_NAME')\\\") 注解",
    "• 5 种角色：CUSTOMER、STAFF、STORE_MANAGER、HQ_OPERATOR、CAT_CARETAKER",
    "• 门店级权限：STAFF/STORE_MANAGER/CAT_CARETAKER 只能操作所属门店数据（Service 层校验）",
    "• 数据归属校验：顾客只能操作自己的预约/订单/退款（通过 AuthPrincipal.userId 比对）",
    "",
    "加密：",
    "• 密码存储：BCrypt 单向哈希（Salt + 多轮加密）",
    "• 通信加密：生产环境使用 HTTPS（课程演示可用 HTTP）",
    "• 敏感配置：JWT Secret、DB 密码、AI API Key 均通过环境变量注入，不硬编码",
    "",
    "审计：",
    "• 关键操作记录：用户创建、角色变更、退款审批、店长分配",
    "• 日志包含：操作人 userId、操作时间、操作类型、目标资源 ID、结果",
    "",
    "合规：",
    "• OWASP Top 10 自查覆盖：注入（MyBatis-Plus 参数化查询防 SQL 注入）、失效鉴权（JWT 校验 + RBAC）、敏感数据暴露（脱敏 + 环境变量）、XML 外部实体（不处理 XML）、失效访问控制（@PreAuthorize + 门店归属校验）、安全配置错误（启动时校验 JWT_SECRET 长度）、XSS（前端框架默认转义）、不安全反序列化（仅使用 JSON）、使用已知漏洞组件（依赖版本锁定）、日志监控不足（traceId + 结构化日志）"
])

# ── Save ──
doc.save(DST)
print(f"D-03 saved to {DST}")
