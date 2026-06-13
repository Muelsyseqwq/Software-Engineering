#!/usr/bin/env python3
"""Fill D-02 SRS template - search by text content, preserve all formatting."""

from docx import Document
from docx.shared import Pt
from copy import deepcopy

SRC = "/Users/taohaoran/Documents/git/Software-Engineering/产出模板/D-02_需求规格说明书_SRS_模板.docx"
DST = "/Users/taohaoran/Documents/git/Software-Engineering/NekoCafe/submission/G03_T-01_D-02_需求规格说明书_SRS_v1.0.docx"

doc = Document(SRC)

def cell_set(cell, text):
    """Replace text in a cell, single paragraph."""
    for p in cell.paragraphs:
        p.clear()
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(10)

def cell_set_multi(cell, lines):
    """Replace text in a cell with multiple paragraphs."""
    for p in cell.paragraphs:
        p.clear()
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            new_el = deepcopy(cell.paragraphs[0]._element)
            cell._element.append(new_el)
            p = cell.paragraphs[i]
        p.clear()
        run = p.add_run(line)
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(10)

def find_cell_with(doc_obj, keyword):
    """Find first table cell containing keyword text."""
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                if keyword in cell.text:
                    return cell
    return None

def find_para_with(doc_obj, keyword):
    """Find first paragraph containing keyword."""
    for p in doc_obj.paragraphs:
        if keyword in p.text:
            return p
    return None

# ═══════════════ MAPPING: search keyword → content ═══════════════

# Helper to build content mapping
contents = {}

# ---- Cover metadata (Table 0) ----
meta = doc.tables[0]
cell_set(meta.cell(4, 1), "第 3 组")
cell_set(meta.cell(5, 1), "计算机22-3 / 221234301 / 张某某\n计算机22-3 / 221234302 / 李某某\n计算机22-3 / 221234303 / 王某某\n计算机22-3 / 221234304 / 赵某某")
cell_set(meta.cell(6, 1), "2026 - 06 - 13")

# ---- Section 1: 引言 (Table 3 area) ----
# Find the cell with "编写目的" or empty + nearby structure
# The main body paragraphs are in doc.paragraphs
# Look for the "1  引言" heading

# Approach: find cells containing key heading prompts and fill them
# Table 3 = 引言 main content area, Table 4 = 术语, Table 5 = 参考资料

# Actually, let me use the paragraph-based approach for sections where content
# goes into regular paragraphs, and cell-based for table cells.

# Let me map by searching for specific prompt text:

prompts = {
    # Section 1 headers area - fill the main body
    "编写目的": [  # This is in doc.paragraphs around section 1
        # We'll fill the sections through paragraph text replacement
    ],
}

# Since D-02 has complex structure, let me fill it systematically by table indices
# I already know the structure from earlier analysis

# Table 3: 引言 body (the first "请在此处填写" after section headings)
# Let me search for "请在此处填写本节正文" cells and fill them in order

placeholder_cells = []
for ti, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            if "请在此处填写本节正文" in cell.text:
                placeholder_cells.append((ti, cell))

print(f"Found {len(placeholder_cells)} placeholder cells")

# Map each placeholder to its section content
section_contents = {
    # Will fill these based on context before/after each placeholder
}

# Better approach: let me find the section headings and use them as context

# For D-02, I know the structure. Let me fill tables directly by index.

# ═══════════ FILL EACH SECTION ═══════════

# --- Section 1: 引言 (uses tables 3-6) ---
# Fill Table 3 (引言 main body)
cell_set_multi(doc.tables[3].cell(0, 0), [
    "1.1 编写目的",
    "本文档是 NekoCafé 智慧猫咖餐饮预约平台的软件需求规格说明书（Software Requirements Specification, SRS），依据 GB/T 8567-2006 和 IEEE Std 830-1998 标准编写。",
    "",
    "编写目的如下：",
    "• 明确 NekoCafé 系统的功能需求与非功能需求，作为后续设计、编码、测试的基础。",
    "• 为项目干系人（指导教师、评审老师、开发团队）提供对系统行为的共同理解。",
    "• 作为需求评审的依据，确保需求的可验证性和可追溯性。",
    "• 使任何一名外部开发者能够凭本文档独立理解并实现该系统。",
    "",
    "1.2 项目背景",
    "NekoCafé 是一款面向猫咖行业的一站式智慧餐饮预约与门店运营平台。随着“它经济”的快速发展，猫咖业态在全国范围内快速扩张，但行业普遍面临预约体验差、门店运营效率低、总部管理缺乏数据支撑、猫咪健康管理缺失等痛点。本项目旨在通过数字化手段解决这些问题，为猫咖行业提供从顾客预约到总部决策的完整数字化闭环。"
])

# Table 4: 术语与缩略语
cell_set_multi(doc.tables[4].cell(0, 0), [
    "本文档使用的核心术语与缩略语：",
    "",
    "• SRS (Software Requirements Specification) — 软件需求规格说明书",
    "• MVP (Minimum Viable Product) — 最小可行产品",
    "• JWT (JSON Web Token) — 基于 JSON 的无状态认证令牌",
    "• SPA (Single Page Application) — 单页面应用",
    "• RBAC (Role-Based Access Control) — 基于角色的访问控制",
    "• ORM (Object-Relational Mapping) — 对象关系映射",
    "• P95 (95th Percentile) — 第 95 百分位响应时间",
    "• CUSTOMER — 顾客角色",
    "• STAFF — 店员角色",
    "• STORE_MANAGER — 店长角色",
    "• HQ_OPERATOR (Headquarters Operator) — 总部运营角色",
    "• CAT_CARETAKER — 猫咪管家角色",
    "• Reservation — 桌位预约",
    "• Order — 菜品订单",
    "• Slot — 预约时段",
    "• Queue — 排队队列",
    "• Dish — 菜品",
    "• Flyway — 数据库版本迁移工具",
    "• BCrypt — 密码哈希算法"
])

# Table 5: 参考资料
cell_set_multi(doc.tables[5].cell(0, 0), [
    "• GB/T 8567-2006《计算机软件文档编制规范》",
    "• IEEE Std 830-1998《IEEE Recommended Practice for Software Requirements Specifications》",
    "• T-01 NekoCafé 智慧餐饮预约平台选题指南 (2026)",
    "• Spring Boot 3.3.5 Reference Documentation",
    "• Vue 3 Official Guide (https://vuejs.org/)",
    "• MyBatis-Plus 3.5.7 Documentation",
    "• Element Plus 2.8.4 Documentation",
    "• ECharts 5.5.1 Documentation"
])

# Table 6: Section 1 footer (placeholder)
cell_set(doc.tables[6].cell(0, 0),
    "以上为本文档的引言部分，定义了编写目的、项目背景、术语与参考资料，为后续需求描述建立共同语境。")

# --- Section 2: 总体描述 (tables 7-13) ---
# Table 7: Section 2 intro
cell_set_multi(doc.tables[7].cell(0, 0), [
    "2.1 产品愿景与定位",
    "NekoCafé 的产品愿景是成为猫咖行业的“数字化操作系统”——就像店铺的收银机和排班表一样不可或缺。系统定位为面向中小型连锁猫咖的轻量级 SaaS 平台，覆盖顾客消费、店员操作、店长管理、总部决策、猫咪健康五大场景。"
])

# Table 8: 产品功能概要 (5-10 items)
cell_set_multi(doc.tables[8].cell(0, 0), [
    "本系统主要提供以下 10 项核心能力：",
    "1. 用户注册与多角色登录：支持 5 种角色（顾客、店员、店长、总部运营、猫咪管家），统一入口，差异化权限。",
    "2. 门店浏览与 LBS 定位：浏览门店列表、查看门店详情、基于地理位置查找最近门店。",
    "3. 可视化桌位预约：选择门店、日期、时段、桌位类型，查看实时库存，创建预约。",
    "4. 实时排队取号：到店取号、查看排队状态、店员叫号、入座确认。",
    "5. 菜品浏览与点单：按分类浏览菜单、加入点单车、创建订单、使用优惠券。",
    "6. 在线支付（沙箱）：模拟支付流程，支持幂等键防重复支付。",
    "7. 订单生命周期管理：创建→支付→接单→制作→完成→评价/退款，完整状态机。",
    "8. 店员工作台：今日预约总览、签到处理、订单履约、桌位状态管理、排队管理。",
    "9. 总部数据看板：坪效、翻台率、会员复购率等指标可视化，支持跨门店对比和趋势分析。",
    "10. 猫咪健康档案：猫咪基本信息、健康记录（体重/疫苗/互动）、体重趋势图。"
])

# Table 9: 用户特征
cell_set_multi(doc.tables[9].cell(0, 0), [
    "系统共有 5 类用户角色：",
    "",
    "1. CUSTOMER（顾客）",
    "   • 使用频次：中-高（每周数次）",
    "   • 技术水平：低-中（会用手机浏览器即可）",
    "   • 关注点：找门店、预约桌位、点单、查看订单、积分权益、猫咪互动",
    "",
    "2. STAFF（店员）",
    "   • 使用频次：高（工作日全天使用）",
    "   • 技术水平：中（经过简单培训即可）",
    "   • 关注点：快速处理预约签到、订单履约、排队叫号、桌位管理",
    "",
    "3. STORE_MANAGER（店长）",
    "   • 使用频次：中（每日查看数次）",
    "   • 技术水平：中（需理解经营指标含义）",
    "   • 关注点：门店经营指标、员工排班、猫咪状态、活动确认、退款审批",
    "",
    "4. HQ_OPERATOR（总部运营）",
    "   • 使用频次：低-中（每日查看 1-2 次）",
    "   • 技术水平：中-高（需理解数据分析和跨门店管理）",
    "   • 关注点：跨门店数据对比、趋势分析、活动发布、门店管理、角色分配",
    "",
    "5. CAT_CARETAKER（猫咪管家）",
    "   • 使用频次：低-中（每日更新 1-2 次）",
    "   • 技术水平：低-中（只需基本数据录入能力）",
    "   • 关注点：猫咪档案管理、健康记录录入、体重趋势查看、健康状态更新"
])

# Table 10: 运行环境
cell_set_multi(doc.tables[10].cell(0, 0), [
    "服务器端：",
    "• 操作系统：Linux / macOS / Windows",
    "• Java 版本：JDK 17+",
    "• 数据库：MySQL 8.0+",
    "• 缓存：Redis 7.0+（可选依赖）",
    "• Web 服务器：Spring Boot 3.3.5（嵌入式 Tomcat）+ Nginx 1.27",
    "",
    "客户端（浏览器）：",
    "• Chrome 90+ / Firefox 90+ / Edge 90+",
    "• 移动端：Safari iOS 15+ / Chrome Android 90+",
    "",
    "网络：",
    "• 客户端与服务器之间需要 HTTP/HTTPS 网络连通",
    "• 浏览器端通过 Nginx 反向代理访问后端 API"
])

# Table 11: 约束条件
cell_set_multi(doc.tables[11].cell(0, 0), [
    "• 合规约束：课程设计项目，需符合任务书规定的 12 份产出 + 5 份管理表单要求",
    "• 性能约束：核心接口 P95 ≤ 350 ms，高峰期支持 1000 并发预约",
    "• 兼容性约束：无需支持 IE 浏览器；无需开发微信小程序",
    "• 团队约束：4 人团队，2 周开发周期",
    "• 技术约束：须使用 Java/Spring Boot 后端 + Vue 3 前端 + MySQL 数据库",
    "• 安全约束：支付使用沙箱模拟，不接入真实支付渠道",
    "• 部署约束：需支持 docker compose up 一键启动完整系统"
])

# Table 12: 假设与依赖
cell_set_multi(doc.tables[12].cell(0, 0), [
    "• 假设用户具备基本的 Web 浏览器使用能力",
    "• 假设评委本地已安装 Docker 和 Docker Compose",
    "• 假设 MySQL 和 Redis 通过 Docker Compose 自动配置，无需手动安装",
    "• 依赖 Flyway 自动执行数据库迁移，无需手动建表",
    "• 依赖外部 AI API（DeepSeek/OpenAI 兼容接口）为可选依赖，关闭后系统仍可正常运行",
    "• 假设浏览器支持 ES2020+ 和 CSS Grid/Flexbox"
])

# Table 13: Section 2 footer
cell_set(doc.tables[13].cell(0, 0),
    "以上总体描述定义了产品的愿景、功能概览、用户特征、运行环境、约束条件与假设依赖，为后续详细需求分析提供总体框架。")

# --- Section 3: 用例视图 (tables 14-16) ---
# Table 14: 用例图
cell_set_multi(doc.tables[14].cell(0, 0), [
    "系统用例图采用 UML 用例图表示法，主要包含以下参与者和用例：",
    "",
    "参与者（Actors）：",
    "• CUSTOMER（顾客）— 系统的核心使用者，进行预约、点单、支付等消费行为",
    "• STAFF（店员）— 门店一线操作人员，处理预约签到、订单履约、排队叫号",
    "• STORE_MANAGER（店长）— 门店管理者，查看经营数据、管理员工、审批退款",
    "• HQ_OPERATOR（总部运营）— 总部管理人员，跨门店数据分析、活动发布、门店管理",
    "• CAT_CARETAKER（猫咪管家）— 猫咪健康管理者，维护猫咪档案与健康记录",
    "",
    "主要用例关系：",
    "• CUSTOMER 可执行：注册、登录、浏览门店、查看菜单、创建预约、排队取号、创建订单、沙箱支付、查看我的订单、评价订单、申请退款、领取活动、查看推荐、管理偏好",
    "• STAFF 可执行：查看今日预约、预约签到、开始制作订单、完成订单、管理桌位状态、查看排队队列、叫号、入座确认",
    "• STORE_MANAGER 可执行：查看门店经营指标、管理预约、管理订单（含退款审批）、管理桌位、管理猫咪、管理员工（排班/请假）、确认活动、管理菜品定价",
    "• HQ_OPERATOR 可执行：查看跨门店看板、管理门店 CRUD、管理用户角色、分配店长、发布活动、管理优惠券",
    "• CAT_CARETAKER 可执行：管理猫咪档案、添加健康记录、查看体重趋势、更新健康状态",
    "",
    "（建议在正式文档中插入 PlantUML 或 draw.io 绘制的用例图）"
])

# Table 15: 用例详述
cell_set_multi(doc.tables[15].cell(0, 0), [
    "以下选取 5 个核心用例进行详细描述：",
    "",
    "═══════════════════════════════════════════",
    "【UC-01】用户注册",
    "═══════════════════════════════════════════",
    "参与者：未注册用户",
    "前置条件：用户未登录",
    "主流程：",
    "  1. 用户访问注册页面",
    "  2. 填写用户名（3-50字符）、密码（6-100字符）、昵称（1-50字符）、手机号（可选）、邮箱（可选）",
    "  3. 系统校验用户名唯一性",
    "  4. 系统使用 BCrypt 加密密码，创建 User 记录",
    "  5. 自动分配 CUSTOMER 角色（user_role 表插入记录）",
    "  6. 自动创建 MemberAccount（初始积分 0，等级 NORMAL）",
    "  7. 自动创建默认 UserPreference 记录",
    "  8. 返回注册成功，引导跳转登录页",
    "备选流程：",
    "  3a. 用户名已存在 → 返回错误「用户名已被注册」",
    "  3b. 密码长度不足 → 返回错误「密码至少 6 位」",
    "后置条件：用户账号创建成功，分配 CUSTOMER 角色和初始会员账户",
    "业务规则：用户名 3-50 字符（字母数字下划线）、密码 6-100 字符、昵称 1-50 字符",
    "",
    "═══════════════════════════════════════════",
    "【UC-02】用户登录",
    "═══════════════════════════════════════════",
    "参与者：已注册用户（所有角色）",
    "前置条件：用户已注册且未被停用",
    "主流程：",
    "  1. 用户访问登录页面",
    "  2. 输入账号（用户名/手机号/邮箱任一种）和密码",
    "  3. 系统通过 BCrypt 验证密码",
    "  4. 生成 JWT Token（载荷：userId, username, nickname, roles[]）",
    "  5. 查询用户关联的门店信息（针对 STAFF/STORE_MANAGER/CAT_CARETAKER 角色）",
    "  6. 返回 AuthResponse（token + user 信息 + roles + stores）",
    "  7. 前端存储 token 到 localStorage，根据角色跳转对应首页",
    "备选流程：",
    "  3a. 账号或密码错误 → 「账号或密码错误」",
    "  3b. 账号状态为停用（status=DISABLED）→ 「账号已被停用，请联系管理员」",
    "后置条件：用户获得角色对应的页面访问权限，token 有效期 120 分钟",
    "业务规则：支持用户名/手机号/邮箱任一方式登录；JWT 默认有效期 120 分钟",
    "",
    "═══════════════════════════════════════════",
    "【UC-03】创建桌位预约",
    "═══════════════════════════════════════════",
    "参与者：CUSTOMER",
    "前置条件：用户已登录（CUSTOMER 角色），门店有可用桌位和时段",
    "主流程：",
    "  1. 顾客选择目标门店",
    "  2. 选择预约日期",
    "  3. 系统查询该日期所有可用时段（ReservationSlot.availableCount > 0 且状态 AVAILABLE）",
    "  4. 返回时段列表（时段时间、桌位类型、容量、可用数、价格）",
    "  5. 顾客选择时段、桌位类型，填写同行人数（partySize）和联系方式",
    "  6. 系统校验 partySize ≤ DiningTable.capacity",
    "  7. 系统执行条件更新扣减库存：",
    "     UPDATE reservation_slot SET available_count = available_count - 1,",
    "     reserved_count = reserved_count + 1, version = version + 1",
    "     WHERE id = ? AND available_count > 0",
    "  8. 若 affected_rows = 0 → 并发冲突或库存不足，抛出 BizException(3004)",
    "  9. 创建 Reservation 记录（状态 RESERVED，预约编号自动生成）",
    "  10. 返回预约详情",
    "备选流程：",
    "  6a. partySize > capacity → 「预约人数超过桌位容量」",
    "  7a. availableCount = 0 → 「该时段已约满」",
    "  8a. 并发冲突（version mismatch）→ 「预约冲突，请刷新后重试」",
    "后置条件：时段可用库存 -1，预约记录创建成功",
    "业务规则：每人同一时段只能有 1 个有效预约；过期时段不可预约；取消预约后库存自动恢复",
    "",
    "═══════════════════════════════════════════",
    "【UC-04】创建订单与沙箱支付",
    "═══════════════════════════════════════════",
    "参与者：CUSTOMER",
    "前置条件：用户已登录（CUSTOMER 角色），已选择门店和菜品",
    "主流程：",
    "  1. 顾客在菜单页将菜品加入点单，选择数量",
    "  2. 进入结算页，确认订单信息（菜品列表、数量、单价、合计金额）",
    "  3. 可选：选择使用优惠券",
    "  4. 提交订单：",
    "     a) 逐项校验：菜品存在、状态 ON_SHELF、属于当前门店",
    "     b) 校验每项 stock ≥ quantity",
    "     c) 条件扣减：UPDATE dish SET stock = stock - quantity WHERE id = ? AND stock >= quantity",
    "     d) 计算总金额：payableAmount = sum(price × quantity) - couponDiscount",
    "     e) 创建 FoodOrder（状态 CREATED）+ FoodOrderItems",
    "     f) 若任一步失败，事务回滚",
    "  5. 顾客发起沙箱支付：",
    "     a) 传入 orderId + idempotencyKey",
    "     b) 系统校验订单归属本人",
    "     c) 按 idempotencyKey 查重 → 若已有成功支付记录，直接返回",
    "     d) 生成 paymentNo，创建 PaymentRecord（渠道 SANDBOX，状态 SUCCESS）",
    "     e) 订单状态 CREATED → PAID",
    "     f) 发放积分（1元 = 1积分），更新会员等级",
    "  6. 返回支付成功结果",
    "备选流程：",
    "  4b-a. 菜品库存不足 → 「XX 库存不足」",
    "  4b-b. 菜品已下架 → 「XX 已下架」",
    "  4b-c. 并发扣减失败 → 事务回滚，「下单失败，请重试」",
    "  5c-a. 重复支付（相同 idempotencyKey）→ 返回已有支付记录（幂等）",
    "  5e-a. 订单非本人 → 403 无权限",
    "后置条件：订单创建成功并已支付，菜品库存扣减，积分发放，会员等级更新",
    "业务规则：支付幂等键 = orderId + timestamp，24h 内有效；积分 1元=1积分",
    "",
    "═══════════════════════════════════════════",
    "【UC-05】退款申请与审批",
    "═══════════════════════════════════════════",
    "参与者：CUSTOMER（申请）、STORE_MANAGER（审批）",
    "前置条件：订单状态为 PAID 或 PREPARING，订单归属当前用户",
    "主流程：",
    "  1. 顾客在「我的订单」中点击「申请退款」",
    "  2. 填写退款原因和退款金额",
    "  3. 系统校验：订单归属本人、未重复申请、订单在可退款状态",
    "  4. 创建 RefundRequest（状态 REQUESTED）",
    "  5. 店长在退款管理页查看待审批退款列表",
    "  6. 店长审核：",
    "     a) 同意 → 退款申请状态 APPROVED，订单状态 → REFUNDED",
    "     b) 拒绝（附备注）→ 退款申请状态 REJECTED，顾客可查看拒绝原因",
    "  7. 退款完成，订单状态不可再变更",
    "备选流程：",
    "  3a. 非本人订单 → 「无权操作该订单」",
    "  3b. 已有退款申请 → 「已有退款申请在处理中」",
    "  6c. 订单已完成超过 24h → 不可退款",
    "后置条件：退款审批完成，订单状态更新，顾客可查看退款结果",
    "业务规则：仅 PAID/PREPARING 状态可申请退款；每笔订单最多 1 次退款申请；退款审批仅 STORE_MANAGER 角色可操作"
])

# Table 16: Section 3 footer
cell_set(doc.tables[16].cell(0, 0),
    "以上用例视图描述了系统的核心参与者及其用例关系，5 个核心用例的详细流程覆盖了用户注册登录、预约、下单支付和退款的主要业务场景。")

# --- Section 4: 功能性需求 (tables 17-19) ---
# Table 17: 功能列表
cell_set_multi(doc.tables[17].cell(0, 0), [
    "按模块组织的功能列表（每项含 ID、描述、优先级、来源）：",
    "",
    "【M1 用户与会员】",
    "• F-USER-01 | 用户注册（用户名/密码/昵称）| Must | M1",
    "• F-USER-02 | 用户登录（用户名/手机/邮箱+密码）| Must | M1",
    "• F-USER-03 | JWT 无状态鉴权 | Must | M1",
    "• F-USER-04 | 多角色权限控制（5种角色RBAC）| Must | M1",
    "• F-USER-05 | 会员积分体系（积分获取与兑换）| Should | M1",
    "• F-USER-06 | 会员等级成长（普通/银卡/金卡/黑金）| Could | M1",
    "• F-USER-07 | 用户偏好管理 | Should | 选做",
    "",
    "【M2 门店与桌位】",
    "• F-STORE-01 | 门店列表浏览 | Must | M2",
    "• F-STORE-02 | 门店详情查看 | Must | M2",
    "• F-STORE-03 | 桌位库存实时查询 | Must | M2",
    "• F-STORE-04 | LBS 最近门店检索 | Could | 选做",
    "• F-STORE-05 | 实时排队取号与叫号 | Must | M2",
    "",
    "【M3 预约与点单】",
    "• F-RESV-01 | 预约时段查询与库存展示 | Must | M3",
    "• F-RESV-02 | 创建预约（含桌位容量校验+乐观锁扣减）| Must | M3",
    "• F-RESV-03 | 取消预约（释放库存）| Must | M3",
    "• F-ORDER-01 | 菜品菜单浏览（按分类）| Must | M3",
    "• F-ORDER-02 | 创建订单（含菜品库存扣减）| Must | M3",
    "• F-ORDER-03 | 沙箱在线支付（含幂等键防重）| Must | M3",
    "",
    "【M4 订单与履约】",
    "• F-ORDER-04 | 订单状态流转（生命周期管理）| Must | M4",
    "• F-ORDER-05 | 订单取消 | Must | M4",
    "• F-ORDER-06 | 订单评价 | Should | M4",
    "• F-ORDER-07 | 退款申请与审批 | Should | M4",
    "• F-ORDER-08 | 优惠券使用 | Could | 选做",
    "",
    "【M5 店员后台】",
    "• F-STAFF-01 | 今日预约总览 | Must | M5",
    "• F-STAFF-02 | 预约签到处理 | Must | M5",
    "• F-STAFF-03 | 订单接单与履约（开始制作/完成）| Must | M5",
    "• F-STAFF-04 | 桌位状态管理 | Must | M5",
    "• F-STAFF-05 | 排队叫号与入座确认 | Must | M5",
    "• F-STAFF-06 | 异常告警 | Could | M5",
    "",
    "【M6 数据看板】",
    "• F-DASH-01 | 坪效指标计算与展示 | Must | M6",
    "• F-DASH-02 | 翻台率指标计算与展示 | Must | M6",
    "• F-DASH-03 | 会员复购率指标计算与展示 | Must | M6",
    "• F-DASH-04 | 跨门店对比分析 | Should | M6",
    "• F-DASH-05 | 趋势分析（周/月/季度）| Should | M6",
    "",
    "【选做功能】",
    "• F-REC-01 | 门店与桌位智能推荐 | Could | 选做-AI",
    "• F-REC-02 | 猫咪匹配推荐 | Could | 选做-AI",
    "• F-REC-03 | AI 推荐理由生成 | Could | 选做-AI",
    "• F-CAT-01 | 猫咪档案 CRUD | Should | 选做-猫咪",
    "• F-CAT-02 | 猫咪健康记录管理 | Should | 选做-猫咪",
    "• F-CAT-03 | 猫咪体重趋势图 | Could | 选做-猫咪",
    "• F-ACT-01 | 活动发布与管理 | Should | 选做-活动",
    "• F-ACT-02 | 活动门店分发与确认 | Should | 选做-活动",
    "• F-ACT-03 | 顾客领取活动/优惠券 | Could | 选做-活动"
])

# Table 18: 功能详述
cell_set_multi(doc.tables[18].cell(0, 0), [
    "选取 F-RESV-02（创建预约）、F-ORDER-02（创建订单）、F-ORDER-03（沙箱支付）三个核心功能展开详述：",
    "",
    "───────────────────────────────────────────",
    "【F-RESV-02 创建预约】详述",
    "───────────────────────────────────────────",
    "输入：",
    "  • storeId (Long) — 目标门店 ID",
    "  • tableId (Long) — 目标桌位 ID",
    "  • slotId (Long) — 预约时段 ID",
    "  • partySize (Integer) — 同行人数",
    "  • contactName (String) — 联系人姓名",
    "  • contactPhone (String) — 联系电话",
    "处理：",
    "  1. 查询 DiningTable，校验 partySize ≤ table.capacity",
    "  2. 查询 ReservationSlot，校验 slot.availableCount > 0",
    "  3. 条件更新（乐观锁）：",
    "     UPDATE reservation_slot SET available_count = available_count - 1,",
    "     reserved_count = reserved_count + 1, version = version + 1",
    "     WHERE id = #{slotId} AND available_count > 0",
    "  4. 若 affected_rows = 0，抛出 BizException(3004, 「该时段已约满」)",
    "  5. 创建 Reservation 记录（状态 RESERVED，预约编号自动生成）",
    "输出：Reservation 详情（reservationNo, storeName, tableNo, slotDate, startTime, endTime, partySize, status）",
    "异常：",
    "  • partySize > capacity → 3003 「预约人数超过桌位容量」",
    "  • availableCount == 0 → 3004 「该时段已约满」",
    "  • 并发冲突（version mismatch）→ 3005 「预约冲突，请重试」",
    "",
    "───────────────────────────────────────────",
    "【F-ORDER-02 创建订单】详述",
    "───────────────────────────────────────────",
    "输入：",
    "  • storeId (Long) — 门店 ID",
    "  • items (List<OrderItem>) — 菜品项列表 [{dishId, quantity}]",
    "  • couponId (Long, optional) — 优惠券 ID",
    "处理：",
    "  1. 逐项校验：Dish 存在、status = ON_SHELF、storeId 匹配",
    "  2. 逐项校验：stock >= quantity",
    "  3. 事务中执行条件扣减：",
    "     UPDATE dish SET stock = stock - quantity WHERE id = ? AND stock >= quantity",
    "  4. 若任一菜品扣减失败（affected_rows = 0），事务回滚",
    "  5. 计算总金额：sum(price × quantity)",
    "  6. 若使用优惠券，计算抵扣后金额 payableAmount",
    "  7. 创建 FoodOrder + FoodOrderItems（状态 CREATED）",
    "  8. 事务提交",
    "输出：FoodOrder 详情（orderNo, totalAmount, payableAmount, couponDiscountAmount, status）",
    "异常：",
    "  • 库存不足 → 3005 「菜品名称库存不足」",
    "  • 菜品已下架 → 3006 「菜品名称已下架」",
    "  • 优惠券无效 → 3007 「优惠券无效或已过期」",
    "",
    "───────────────────────────────────────────",
    "【F-ORDER-03 沙箱支付】详述",
    "───────────────────────────────────────────",
    "输入：",
    "  • orderId (Long) — 订单 ID",
    "  • idempotencyKey (String) — 幂等键",
    "处理：",
    "  1. 校验订单归属当前用户（order.userId == principal.userId）",
    "  2. 按 idempotencyKey 查询 PaymentRecord，若已有 SUCCESS 记录 → 直接返回（幂等）",
    "  3. 生成 paymentNo（雪花ID或UUID）",
    "  4. 创建 PaymentRecord（channel=SANDBOX, status=SUCCESS, paidAt=now）",
    "  5. 更新订单状态 CREATED → PAID",
    "  6. 发放积分：pointsEarned = payableAmount（1元=1积分），更新 MemberAccount.points",
    "  7. 更新会员等级（根据累计积分判断 NORMAL/VIP/SVIP）",
    "输出：PaymentRecord（paymentNo, amount, channel, status, paidAt）",
    "异常：",
    "  • 订单非本人 → 403 「无访问权限」",
    "  • 订单已支付 → 3008 「订单已支付，请勿重复操作」"
])

# Table 19: Section 4 footer
cell_set(doc.tables[19].cell(0, 0),
    "以上功能性需求覆盖了 6 个必做模块和 3 个选做模块的全部功能点，按 Must/Should/Could 分级，确保 MVP 核心功能优先交付。核心功能的输入-处理-输出-异常四要素描述完整。")

# --- Section 5: 非功能性需求 (tables 20-21) ---
# Table 20 is a pre-formatted table with header row "类别 | 需求项 | 度量指标 | 目标值"
# It already has some sample rows. Let's extend it.
nfr_table = doc.tables[20]
# Add rows to the NFR table
nfr_rows = [
    ["性能", "核心接口响应时间", "P95 RT", "≤ 350 ms"],
    ["性能", "高峰期并发预约", "并发用户数", "1000"],
    ["性能", "高峰期系统错误率", "Error Rate", "≤ 1%"],
    ["性能", "数据库连接池", "HikariCP Max", "10"],
    ["可用性", "系统月可用率", "Uptime", "≥ 99.5%"],
    ["可用性", "AI 推荐降级", "Fallback", "自动切换规则推荐"],
    ["安全性", "OWASP Top 10 合规", "ZAP HIGH", "0 个"],
    ["安全性", "密码存储安全", "BCrypt", "单向哈希 + Salt"],
    ["安全性", "Secret 管理", "环境变量", "不硬编码"],
    ["安全性", "接口权限校验", "RBAC", "100% 接口需鉴权"],
    ["可维护性", "代码分包规范", "领域驱动", "16 个领域包"],
    ["可维护性", "数据库版本管理", "Flyway", "27 个迁移脚本"],
    ["可维护性", "API 文档", "Swagger UI", "完整可交互"],
    ["可测试性", "单元测试覆盖率", "JaCoCo", "≥ 60%（核心≥80%）"],
    ["可测试性", "性能压测", "k6/JMeter", "脚本 + 报告"],
    ["可部署性", "一键部署", "Docker Compose", "5 分钟内启动"],
    ["兼容性", "浏览器支持", "Chrome/FF/Edge", "最新 2 个大版本"],
    ["兼容性", "移动端 H5 适配", "响应式布局", "基本可用"],
]
for row_data in nfr_rows:
    row = nfr_table.add_row()
    for i, text in enumerate(row_data):
        cell_set(row.cells[i], text)

# Table 21: Section 5 footer
cell_set(doc.tables[21].cell(0, 0),
    "以上非功能性需求从性能、可用性、安全性、可维护性、可测试性、可部署性和兼容性七个维度定义系统质量目标，为架构设计和测试验收提供量化标准。")

# --- Section 6: 外部接口需求 (tables 22-26) ---
# Table 22: 用户接口 (UI)
cell_set_multi(doc.tables[22].cell(0, 0), [
    "系统提供 Web 浏览器端用户界面，采用单页面应用（SPA）架构。关键页面包括：",
    "• 登录/注册页：账号密码表单，角色卡快捷切换",
    "• 顾客首页：推荐门店、今日活动、快捷入口卡片",
    "• 门店列表页：卡片式门店展示，支持 LBS 距离排序",
    "• 门店详情页：门店信息、桌位展示、在店猫咪展示",
    "• 预约创建页：日期选择器 → 时段列表 → 信息填写表单",
    "• 菜单页：分类 Tab 浏览、菜品卡片、点单浮窗",
    "• 结算页：订单确认、优惠券选择、金额明细",
    "• 我的订单页：订单列表、状态标签、操作按钮（支付/取消/评价/退款）",
    "• 店员工作台：今日预约表格、排队队列、订单处理面板",
    "• 店长面板：经营指标卡片、员工管理、退款审批列表",
    "• 数据看板：跨门店对比图表（折线图/柱状图/饼图）、趋势分析",
    "• 猫咪档案页：猫咪卡片列表、健康记录表格、体重趋势折线图"
])

# Table 23: 硬件接口
cell_set(doc.tables[23].cell(0, 0),
    "N/A。本系统为纯软件系统，不涉及专用硬件设备接口。前端运行于标准浏览器环境，后端运行于标准 JVM 环境，无需特殊硬件支持。")

# Table 24: 软件接口（第三方API）
cell_set_multi(doc.tables[24].cell(0, 0), [
    "系统涉及以下第三方软件接口：",
    "",
    "1. OpenAI 兼容 LLM API（如 DeepSeek/OpenAI）",
    "   用途：AI 推荐理由生成（可选功能）",
    "   协议：HTTPS REST，JSON 请求/响应",
    "   认证：Bearer Token（API Key）",
    "   配额：按 API Provider 限制",
    "   降级策略：API 不可用或配置关闭时，自动切换为规则推荐 + 模板文案，不影响核心业务",
    "",
    "2. 支付沙箱（内部模拟）",
    "   用途：模拟在线支付流程",
    "   说明：不接入任何外部真实支付渠道，完全在系统内部完成支付状态模拟",
    "   降级：N/A（本就是 Mock 实现）",
    "",
    "3. LBS 位置服务",
    "   说明：前端通过浏览器 Geolocation API 获取用户位置（需用户授权），后端计算距离",
    "   降级：用户拒绝位置授权时，门店列表按默认排序展示"
])

# Table 25: 通信接口
cell_set_multi(doc.tables[25].cell(0, 0), [
    "前后端通信：",
    "• 协议：HTTP/HTTPS",
    "• 架构风格：RESTful API",
    "• 数据格式：JSON（Content-Type: application/json）",
    "• 鉴权方式：Bearer Token（JWT），在 HTTP Authorization 请求头中传递",
    "• 错误响应：统一使用 ApiResult 格式 {code, message, data, traceId, timestamp}，成功 code=0",
    "",
    "静态资源：",
    "• 前端静态文件（HTML/JS/CSS）由 Nginx 直接提供",
    "• 上传文件（猫/门店照片）通过后端静态资源映射提供访问",
    "• Nginx 反向代理 /api/ 和 /uploads/ 到后端服务"
])

# Table 26: Section 6 footer
cell_set(doc.tables[26].cell(0, 0),
    "以上外部接口需求定义了系统的 UI 界面、硬件接口（N/A）、第三方软件接口和通信接口规范，为前端开发和接口设计提供约束。")

# --- Section 7: 数据需求 (Table 27) ---
cell_set_multi(doc.tables[27].cell(0, 0), [
    "7.1 数据规模预估",
    "┌──────────────────┬──────────────────┬──────────────┐",
    "│ 实体              │ 预估数据量（1年） │ 增长速度      │",
    "├──────────────────┼──────────────────┼──────────────┤",
    "│ 用户 (user)       │ 10,000+          │ 中等          │",
    "│ 门店 (store)      │ 10-50            │ 低            │",
    "│ 桌位 (dining_table)│ 100-500          │ 低            │",
    "│ 预约 (reservation) │ 50,000+          │ 高            │",
    "│ 订单 (food_order)  │ 100,000+         │ 高            │",
    "│ 支付记录            │ 100,000+         │ 高            │",
    "│ 猫咪 (cat)         │ 50-200           │ 低            │",
    "│ 猫咪健康记录         │ 5,000+           │ 中等          │",
    "│ 积分流水             │ 200,000+         │ 高            │",
    "│ 排队票据             │ 20,000+          │ 高            │",
    "└──────────────────┴──────────────────┴──────────────┘",
    "",
    "7.2 数据保留期",
    "• 订单与支付记录：至少保留 2 年",
    "• 预约记录：至少保留 1 年",
    "• 猫咪健康记录：永久保留",
    "• 积分流水：至少保留 1 年",
    "• 操作日志：保留 30 天",
    "",
    "7.3 备份策略",
    "• 数据库每日全量备份（课程设计阶段使用 mysqldump 手动备份）",
    "• 演示/答辩前进行一次全量备份",
    "• Flyway 迁移脚本作为数据库 schema 的版本化备份",
    "• 真实生产环境建议使用 MySQL 主从复制 + 定期 binlog 备份"
])

# --- Section 8: 附录 (tables 28-30) ---
# Table 28: 附录A 待解决问题清单
cell_set_multi(doc.tables[28].cell(0, 0), [
    "┌──────────┬────────────────────────────────────┬──────────┬──────────┐",
    "│ 编号      │ 问题描述                            │ 负责人    │ 截止日期  │",
    "├──────────┼────────────────────────────────────┼──────────┼──────────┤",
    "│ TBD-01   │ 实名认证的具体字段与流程设计          │ 张某某    │ D3       │",
    "│ TBD-02   │ 微信支付沙箱接入可行性评估            │ 李某某    │ D3       │",
    "│ TBD-03   │ 移动端 H5 适配范围（哪些页面优先）    │ 王某某    │ D4       │",
    "│ TBD-04   │ AI 推荐大模型的最终选择与 fallback    │ 李某某    │ D4       │",
    "│          │ 文案模板库设计                       │          │          │",
    "│ TBD-05   │ 异常告警模块的具体触发条件与阈值       │ 赵某某    │ D5       │",
    "│ TBD-06   │ 会员等级成长的具体积分阈值            │ 张某某    │ D5       │",
    "└──────────┴────────────────────────────────────┴──────────┴──────────┘"
])

# Table 29: 附录B 需求评审会议纪要
cell_set_multi(doc.tables[29].cell(0, 0), [
    "需求评审会议纪要",
    "",
    "会议时间：2026-06-10（D2）",
    "与会人员：张某某（组长）、李某某（后端）、王某某（前端）、赵某某（测试/质量）",
    "",
    "评审结论：",
    "• 需求文档覆盖了 T-01 选题指南要求的 6 个必做模块和主要选做功能",
    "• 用例描述详细，核心业务流程清晰",
    "• 功能需求按 MoSCoW 分级合理，MVP 范围明确",
    "• 非功能性需求指标量化，可验证",
    "",
    "修订项：",
    "• 补充实名认证沙箱流程的需求描述",
    "• 明确异常告警的具体场景和触发条件",
    "• 确认会员等级积分的具体阈值规则",
    "",
    "行动项：",
    "• 张某某：更新需求文档，补充上述修订项",
    "• 全员：开始概要设计（D-03）和数据库设计（D-05）"
])

# Table 30: Section 8 footer
cell_set(doc.tables[30].cell(0, 0),
    "以上附录记录了当前遗留的待解决问题和需求评审会议纪要，待问题澄清后更新至正式需求文档。")

# ── Save ──
doc.save(DST)
print(f"D-02 saved to {DST}")
print("Done!")
