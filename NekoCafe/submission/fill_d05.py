#!/usr/bin/env python3
"""Fill D-05 数据库设计说明书 - preserve formatting, use 「」 for quotes."""

from docx import Document
from docx.shared import Pt
from copy import deepcopy
import glob, os

matches = glob.glob("/Users/taohaoran/Documents/git/Software-Engineering/产出模板/D-05_*模板.docx")
SRC = matches[0]
DST = "/Users/taohaoran/Documents/git/Software-Engineering/NekoCafe/submission/G03_T-01_D-05_数据库设计说明书_v1.0.docx"

doc = Document(SRC)

def cs(cell, text):
    for p in cell.paragraphs: p.clear()
    for p in cell.paragraphs[1:]: p._element.getparent().remove(p._element)
    cell.paragraphs[0].text = ""
    r = cell.paragraphs[0].add_run(text)
    r.font.name = "Microsoft YaHei"; r.font.size = Pt(10)

def cm(cell, lines):
    for p in cell.paragraphs: p.clear()
    for p in cell.paragraphs[1:]: p._element.getparent().remove(p._element)
    for i, line in enumerate(lines):
        if i == 0: p = cell.paragraphs[0]
        else:
            new_el = deepcopy(cell.paragraphs[0]._element)
            cell._element.append(new_el)
            p = cell.paragraphs[i]
        p.clear()
        r = p.add_run(line)
        r.font.name = "Microsoft YaHei"; r.font.size = Pt(10)

# ── Cover metadata ──
t = doc.tables[0]
cs(t.cell(4,1), "第 3 组")
cs(t.cell(5,1), "计算机22-3 / 221234301 / 张某某\n计算机22-3 / 221234302 / 李某某\n计算机22-3 / 221234303 / 王某某\n计算机22-3 / 221234304 / 赵某某")
cs(t.cell(6,1), "2026 - 06 - 13")

# ═══════ D-05 TABLE MAP ═══════
# T0: metadata | T1: guidance | T2: cover info | T3: sec 1 heading
# T4: sec 1 引言 (placeholder)
# T5: sec 2 概念模型 ER图 (guided)
# T6: sec 3.1 关系模式 | T7: sec 3.2 范式 (guided)
# T8: sec 3 footer (placeholder)
# T9: sec 4.1 数据字典 | T10: sec 4.2 DDL
# T11: sec 4 footer (placeholder)
# T12: sec 5.1 索引 | T13: sec 5.2 性能预估
# T14: sec 5 footer (placeholder)
# T15: sec 6.1 权限 | T16: sec 6.2 敏感字段 | T17: sec 6.3 审计备份
# T18: sec 6 footer (placeholder)
# T19: sec 7 数据生命周期

# Section 1: 引言 (T4)
cm(doc.tables[4].cell(0,0), [
    "1.1 编写目的",
    "本文档是 NekoCafé 智慧猫咖餐饮预约平台的数据库设计说明书，描述系统的数据组织、存储、访问和管理策略，涵盖概念模型、逻辑模型、物理模型及数据安全与合规。",
    "",
    "1.2 数据库选型与版本",
    "选型：MySQL 8.0 + InnoDB 存储引擎",
    "选型理由：",
    "• 成熟稳定：全球最流行的开源关系型数据库，社区活跃，资料丰富",
    "• 事务支持：InnoDB 支持 ACID 事务、行级锁、MVCC，适合高并发预约/订单场景",
    "• 字符集：utf8mb4 完整支持 Unicode（包括 emoji），满足猫咪名字等特殊字符需求",
    "• Docker 化：官方 Docker 镜像成熟，一键部署，环境一致",
    "• 生态整合：Spring Boot + MyBatis-Plus + Flyway 对 MySQL 支持最佳",
    "• 团队熟悉：团队成员在课程中已有 MySQL 学习和使用经验",
    "",
    "版本管理：使用 Flyway 管理数据库版本，所有 DDL/DML 变更通过 25+ 个迁移脚本执行，确保数据库 schema 的可追溯、可复现。"
])

# Section 2: 概念模型 ER 图 (T5)
cm(doc.tables[5].cell(0,0), [
    "以下按业务领域分层描述实体关系（ER 模型）。",
    "",
    "【全局 ER 图 — 核心实体关系】",
    "",
    "  ┌──────┐     ┌───────────┐     ┌──────┐",
    "  │ user │────<│ user_role │>────│ role │",
    "  └──┬───┘     └───────────┘     └──────┘",
    "     │",
    "     ├── member_account (1:1)",
    "     ├── user_preference (1:N)",
    "     ├── user_store_role (1:N) ─── store",
    "     ├── reservation (1:N) ─── store, dining_table, reservation_slot",
    "     ├── food_order (1:N) ─── store, reservation, dining_table",
    "     ├── review (1:N) ─── food_order",
    "     ├── refund_request (1:N) ─── food_order",
    "     ├── points_transaction (1:N)",
    "     ├── reward_redemption (1:N) ─── reward_catalog",
    "     └── waiting_queue_ticket (1:N) ─── store",
    "",
    "  ┌───────┐     ┌───────────────┐     ┌────────────────┐",
    "  │ store │────<│ dining_table  │────<│reservation_slot│",
    "  └──┬────┘     └───────────────┘     └────────────────┘",
    "     │",
    "     ├── dish_category (1:N) ─── dish",
    "     ├── cat (1:N) ─── cat_health_record",
    "     ├── staff_shift (1:N) ─── user",
    "     ├── waiting_queue_counter (1:1)",
    "     ├── activity_store (1:N) ─── promotion_activity",
    "     ├── store_manager_assignment (1:N) ─── user",
    "     └── store_photo (1:N)",
    "",
    "  ┌──────────────┐",
    "  │ food_order   │───< food_order_item >─── dish",
    "  │              │─── payment_record (1:1)",
    "  └──────────────┘",
    "",
    "  ┌────────────────────┐",
    "  │ promotion_activity │───< activity_store >── store",
    "  │                    │───< reward_catalog",
    "  └────────────────────┘",
    "",
    "【子图 1：预约与排队域】",
    "  user ──< reservation >── dining_table",
    "  user ──< reservation >── reservation_slot >── dining_table",
    "  user ──< waiting_queue_ticket >── waiting_queue_counter >── store",
    "",
    "【子图 2：订单与支付域】",
    "  user ──< food_order >──< food_order_item >── dish >── dish_category >── store",
    "  food_order ── payment_record",
    "  food_order ──< review",
    "  food_order ──< refund_request",
    "  user ──< points_transaction",
    "  user ──< reward_redemption >── reward_catalog >── promotion_activity",
    "",
    "【子图 3：猫咪与员工域】",
    "  store ──< cat >──< cat_health_record",
    "  store ──< user_store_role >── user >── role",
    "  store ──< staff_shift >── user",
    "  store ──< staff_leave_request >── user"
])

# Section 3: 逻辑模型 (T6, T7, T8)
cm(doc.tables[6].cell(0,0), [
    "3.1 关系模式（核心表）",
    "",
    "user(user_id, username, password_hash, nickname, phone, email, status, created_at, updated_at, deleted, version)",
    "  PK: user_id  |  UK: username",
    "",
    "role(role_id, code, name, description)",
    "  PK: role_id  |  UK: code",
    "",
    "user_role(id, user_id, role_id)",
    "  PK: id  |  FK: user_id→user, role_id→role",
    "",
    "member_account(id, user_id, level_code, points, total_spent)",
    "  PK: id  |  FK: user_id→user  |  UK: user_id",
    "",
    "store(store_id, name, city, address, phone, opening_time, closing_time, status, description, business_area, latitude, longitude, cover_url, area_square_meter)",
    "  PK: store_id",
    "",
    "dining_table(table_id, store_id, table_no, capacity, area, status)",
    "  PK: table_id  |  FK: store_id→store",
    "",
    "reservation_slot(slot_id, store_id, table_id, slot_date, start_time, end_time, capacity, reserved_count, available_count, status, version)",
    "  PK: slot_id  |  FK: store_id→store, table_id→dining_table",
    "",
    "reservation(reservation_id, reservation_no, user_id, store_id, table_id, slot_id, party_size, status, contact_name, contact_phone, checked_in_at, created_at, version)",
    "  PK: reservation_id  |  FK: user_id→user, store_id→store, table_id→dining_table, slot_id→reservation_slot",
    "",
    "dish_category(category_id, store_id, name, sort_order, status)",
    "  PK: category_id  |  FK: store_id→store",
    "",
    "dish(dish_id, store_id, category_id, name, price, stock, status, description, image_url)",
    "  PK: dish_id  |  FK: store_id→store, category_id→dish_category",
    "",
    "food_order(order_id, order_no, user_id, store_id, reservation_id, table_id, total_amount, reward_redemption_id, coupon_discount_amount, payable_amount, coupon_name, status, refund_status, handler_id, created_at, version)",
    "  PK: order_id  |  FK: user_id→user, store_id→store",
    "",
    "food_order_item(item_id, order_id, dish_id, dish_name, unit_price, quantity, subtotal)",
    "  PK: item_id  |  FK: order_id→food_order, dish_id→dish",
    "",
    "payment_record(payment_id, payment_no, order_id, idempotency_key, amount, channel, status, paid_at)",
    "  PK: payment_id  |  FK: order_id→food_order  |  UK: idempotency_key",
    "",
    "cat(cat_id, store_id, name, breed, age, weight, gender, personality, interact, health_status, vaccinium, photo_url, description, status)",
    "  PK: cat_id  |  FK: store_id→store",
    "",
    "cat_health_record(record_id, cat_id, weight, temperature, appetite_status, vaccine_status, interaction_note, recorded_at, created_by)",
    "  PK: record_id  |  FK: cat_id→cat",
    "",
    "user_preference(id, user_id, pref_type, pref_value)",
    "  PK: id  |  FK: user_id→user  |  UK: (user_id, pref_type)",
    "",
    "review(review_id, user_id, order_id, rating, content, created_at)",
    "  PK: review_id  |  FK: user_id→user, order_id→food_order  |  UK: order_id (每单一条)",
    "",
    "refund_request(refund_id, user_id, order_id, reason, amount, status, handled_by, handle_note, created_at, handled_at)",
    "  PK: refund_id  |  FK: user_id→user, order_id→food_order",
    "",
    "waiting_queue_counter(counter_id, store_id, current_number, status)",
    "  PK: counter_id  |  FK: store_id→store",
    "",
    "waiting_queue_ticket(ticket_id, store_id, user_id, ticket_number, party_size, status, created_at)",
    "  PK: ticket_id  |  FK: store_id→store, user_id→user",
    "",
    "promotion_activity(activity_id, title, description, start_time, end_time, status, created_by)",
    "  PK: activity_id  |  FK: created_by→user",
    "",
    "activity_store(id, activity_id, store_id, accept_status)",
    "  PK: id  |  FK: activity_id→promotion_activity, store_id→store  |  UK: (activity_id, store_id)",
    "",
    "dashboard_stat(id, store_id, stat_date, revenue, order_count, reservation_count, avg_rating)",
    "  PK: id  |  FK: store_id→store",
    "",
    "points_transaction(id, user_id, amount, type, reference_id, description, created_at)",
    "  PK: id  |  FK: user_id→user",
    "",
    "reward_catalog(reward_id, activity_id, name, points_cost, discount_amount, reward_type, required_level)",
    "  PK: reward_id  |  FK: activity_id→promotion_activity",
    "",
    "reward_redemption(id, user_id, reward_id, activity_id, status, created_at)",
    "  PK: id  |  FK: user_id→user, reward_id→reward_catalog  |  UK: (user_id, activity_id)"
])

cm(doc.tables[7].cell(0,0), [
    "3.2 范式说明与反范式取舍",
    "",
    "范式合规：全局设计达到第三范式（3NF），所有非主键字段完全函数依赖于主键，不存在传递依赖和部分依赖。",
    "",
    "反范式取舍（为提高查询性能的刻意反范式）：",
    "",
    "1. food_order 表中的冗余字段：",
    "   • coupon_name (String)：冗余存储优惠券名称，避免每次查询订单列表时 JOIN coupon 表",
    "   • coupon_discount_amount (Decimal)：冗余存储折扣金额",
    "   • total_amount (Decimal)：冗余存储订单原总金额",
    "   • 理由：订单列表（顾客「我的订单」、店员「订单处理」、店长「订单管理」）是高频查询场景，冗余字段减少 JOIN，提升查询性能",
    "",
    "2. food_order_item 表中的冗余字段：",
    "   • dish_name (String)：冗余存储菜品名称，保留下单时的菜品名称快照",
    "   • unit_price (Decimal)：冗余存储下单时的单价",
    "   • 理由：菜品名称和价格可能被店长修改，订单项需要保留下单时的历史快照，确保订单数据不可篡改",
    "",
    "3. reservation_slot 表中的派生字段：",
    "   • reserved_count (Integer)：派生自 reservation 表的 COUNT 聚合",
    "   • available_count (Integer)：派生计算 = capacity - reserved_count",
    "   • 理由：高并发预约场景下，每次查询可用库存都 COUNT reservation 表代价过高。冗余 available_count 并通过乐观锁维护一致性，显著提升查询性能"
])

cs(doc.tables[8].cell(0,0),
    "以上逻辑模型定义了 25 张核心表的关系模式和范式设计，3 处反范式取舍均有明确的性能优化理由和一致性保障措施。")

# Section 4: 物理模型与 DDL (T9, T10, T11)
cm(doc.tables[9].cell(0,0), [
    "4.1 数据字典（核心表字段说明）",
    "",
    "═══════════════════════════════════════════",
    "表：user（用户表）",
    "═══════════════════════════════════════════",
    "┌─────────────────┬──────────────┬──────┬──────┬──────┬──────────────────┬──────┐",
    "│ 字段             │ 类型          │ 主键  │ 可空  │ 默认  │ 说明              │ 敏感 │",
    "├─────────────────┼──────────────┼──────┼──────┼──────┼──────────────────┼──────┤",
    "│ user_id         │ BIGINT       │ PK   │ N    │ 雪花  │ 用户 ID           │ N    │",
    "│ username        │ VARCHAR(50)  │ UK   │ N    │ -    │ 用户名             │ N    │",
    "│ password_hash   │ VARCHAR(255) │      │ N    │ -    │ BCrypt 密码哈希    │ Y    │",
    "│ nickname        │ VARCHAR(50)  │      │ N    │ -    │ 昵称               │ N    │",
    "│ phone           │ VARCHAR(20)  │      │ Y    │ NULL │ 手机号             │ Y    │",
    "│ email           │ VARCHAR(100) │      │ Y    │ NULL │ 邮箱               │ Y    │",
    "│ status          │ VARCHAR(20)  │      │ N    │ACTIVE│ 状态（ACTIVE/DISABLED）│ N    │",
    "│ created_at      │ DATETIME     │      │ N    │ NOW()│ 创建时间           │ N    │",
    "│ updated_at      │ DATETIME     │      │ N    │ NOW()│ 更新时间           │ N    │",
    "│ deleted         │ TINYINT      │      │ N    │ 0    │ 逻辑删除（1=已删） │ N    │",
    "│ version         │ INT          │      │ N    │ 0    │ 乐观锁版本号       │ N    │",
    "└─────────────────┴──────────────┴──────┴──────┴──────┴──────────────────┴──────┘",
    "",
    "═══════════════════════════════════════════",
    "表：reservation（预约表）",
    "═══════════════════════════════════════════",
    "┌─────────────────┬──────────────┬──────┬──────┬──────┬──────────────────┬──────┐",
    "│ 字段             │ 类型          │ 主键  │ 可空  │ 默认  │ 说明              │ 敏感 │",
    "├─────────────────┼──────────────┼──────┼──────┼──────┼──────────────────┼──────┤",
    "│ reservation_id  │ BIGINT       │ PK   │ N    │ 雪花  │ 预约 ID           │ N    │",
    "│ reservation_no  │ VARCHAR(32)  │ UK   │ N    │ -    │ 预约编号（业务键） │ N    │",
    "│ user_id         │ BIGINT       │      │ N    │ -    │ FK→user           │ N    │",
    "│ store_id        │ BIGINT       │      │ N    │ -    │ FK→store          │ N    │",
    "│ table_id        │ BIGINT       │      │ N    │ -    │ FK→dining_table   │ N    │",
    "│ slot_id         │ BIGINT       │      │ N    │ -    │ FK→reservation_slot│ N    │",
    "│ party_size      │ INT          │      │ N    │ -    │ 同行人数           │ N    │",
    "│ status          │ VARCHAR(20)  │      │ N    │RESVED│ RESERVED/CANCELLED/CHECKED_IN│N│",
    "│ contact_name    │ VARCHAR(50)  │      │ Y    │ NULL │ 联系人姓名         │ Y    │",
    "│ contact_phone   │ VARCHAR(20)  │      │ Y    │ NULL │ 联系电话           │ Y    │",
    "│ checked_in_at   │ DATETIME     │      │ Y    │ NULL │ 签到时间           │ N    │",
    "│ created_at      │ DATETIME     │      │ N    │ NOW()│ 创建时间           │ N    │",
    "│ version         │ INT          │      │ N    │ 0    │ 乐观锁版本号       │ N    │",
    "└─────────────────┴──────────────┴──────┴──────┴──────┴──────────────────┴──────┘",
    "",
    "═══════════════════════════════════════════",
    "表：food_order（订单表）",
    "═══════════════════════════════════════════",
    "┌──────────────────────┬──────────────┬──────┬──────┬──────┬──────────────────┬──────┐",
    "│ 字段                  │ 类型          │ 主键  │ 可空  │ 默认  │ 说明              │ 敏感 │",
    "├──────────────────────┼──────────────┼──────┼──────┼──────┼──────────────────┼──────┤",
    "│ order_id             │ BIGINT       │ PK   │ N    │ 雪花  │ 订单 ID           │ N    │",
    "│ order_no             │ VARCHAR(32)  │ UK   │ N    │ -    │ 订单编号（业务键） │ N    │",
    "│ user_id              │ BIGINT       │      │ N    │ -    │ FK→user           │ N    │",
    "│ store_id             │ BIGINT       │      │ N    │ -    │ FK→store          │ N    │",
    "│ total_amount         │ DECIMAL(10,2)│      │ N    │ -    │ 原总金额（冗余）   │ N    │",
    "│ coupon_discount_amount│DECIMAL(10,2)│     │ Y    │ 0.00 │ 优惠券抵扣金额     │ N    │",
    "│ payable_amount       │ DECIMAL(10,2)│      │ N    │ -    │ 实付金额           │ N    │",
    "│ coupon_name          │ VARCHAR(100) │      │ Y    │ NULL │ 优惠券名称（冗余） │ N    │",
    "│ status               │ VARCHAR(20)  │      │ N    │CREATED│ 订单状态          │ N    │",
    "│ handler_id           │ BIGINT       │      │ Y    │ NULL │ 处理店员 ID        │ N    │",
    "│ created_at           │ DATETIME     │      │ N    │ NOW()│ 创建时间           │ N    │",
    "└──────────────────────┴──────────────┴──────┴──────┴──────┴──────────────────┴──────┘",
    "",
    "其余 22 张表的数据字典详见 Flyway 迁移脚本（V001-V025），此处不逐表展开。"
])

cm(doc.tables[10].cell(0,0), [
    "4.2 DDL 脚本",
    "所有 DDL 脚本通过 Flyway 迁移文件管理，位于 src/main/resources/db/migration/。",
    "迁移文件按版本号 V001-V025 按依赖顺序执行：",
    "",
    "核心建表脚本（V001__init_mysql_schema.sql）包含以下主要 DDL：",
    "",
    "CREATE TABLE user (",
    "    user_id BIGINT PRIMARY KEY,",
    "    username VARCHAR(50) NOT NULL UNIQUE,",
    "    password_hash VARCHAR(255) NOT NULL,",
    "    nickname VARCHAR(50) NOT NULL,",
    "    phone VARCHAR(20),",
    "    email VARCHAR(100),",
    "    status VARCHAR(20) NOT NULL DEFAULT 'ACTIVE',",
    "    created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,",
    "    updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,",
    "    deleted TINYINT NOT NULL DEFAULT 0,",
    "    version INT NOT NULL DEFAULT 0",
    ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;",
    "",
    "CREATE TABLE reservation_slot (",
    "    slot_id BIGINT PRIMARY KEY,",
    "    store_id BIGINT NOT NULL,",
    "    table_id BIGINT NOT NULL,",
    "    slot_date DATE NOT NULL,",
    "    start_time TIME NOT NULL,",
    "    end_time TIME NOT NULL,",
    "    capacity INT NOT NULL DEFAULT 1,",
    "    reserved_count INT NOT NULL DEFAULT 0,",
    "    available_count INT NOT NULL DEFAULT 1,",
    "    status VARCHAR(20) NOT NULL DEFAULT 'AVAILABLE',",
    "    version INT NOT NULL DEFAULT 0,",
    "    INDEX idx_slot_store_date (store_id, slot_date),",
    "    INDEX idx_slot_table (table_id),",
    "    FOREIGN KEY (store_id) REFERENCES store(store_id),",
    "    FOREIGN KEY (table_id) REFERENCES dining_table(table_id)",
    ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;",
    "",
    "CREATE TABLE payment_record (",
    "    payment_id BIGINT PRIMARY KEY,",
    "    payment_no VARCHAR(32) NOT NULL UNIQUE,",
    "    order_id BIGINT NOT NULL,",
    "    idempotency_key VARCHAR(128) NOT NULL UNIQUE,",
    "    amount DECIMAL(10,2) NOT NULL,",
    "    channel VARCHAR(20) NOT NULL DEFAULT 'SANDBOX',",
    "    status VARCHAR(20) NOT NULL DEFAULT 'PENDING',",
    "    paid_at DATETIME,",
    "    FOREIGN KEY (order_id) REFERENCES food_order(order_id)",
    ") ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;",
    "",
    "完整的 25 个迁移脚本清单见 D-12 项目源代码仓库说明。"
])

cs(doc.tables[11].cell(0,0),
    "以上物理模型通过数据字典和 DDL 脚本定义了核心表的完整结构，所有 DDL 通过 Flyway 版本化管理，确保数据库 schema 的可追溯性和可复现性。")

# Section 5: 索引与性能 (T12, T13, T14)
cm(doc.tables[12].cell(0,0), [
    "5.1 索引清单",
    "┌──────────────────┬──────────────────────┬──────────────┬──────────────────────┐",
    "│ 表名              │ 索引字段              │ 索引类型      │ 选择性/说明            │",
    "├──────────────────┼──────────────────────┼──────────────┼──────────────────────┤",
    "│ user             │ username             │ UNIQUE       │ 极高（唯一）           │",
    "│ user             │ phone                │ INDEX        │ 高（登录查询）         │",
    "│ user             │ email                │ INDEX        │ 高（登录查询）         │",
    "│ reservation_slot │ (store_id, slot_date)│ COMPOSITE    │ 中（时段查询核心）      │",
    "│ reservation_slot │ table_id             │ INDEX        │ 中（按桌位过滤）        │",
    "│ reservation      │ user_id              │ INDEX        │ 中（我的预约）          │",
    "│ reservation      │ slot_id              │ INDEX        │ 高（库存关联）          │",
    "│ food_order       │ user_id              │ INDEX        │ 中（我的订单）          │",
    "│ food_order       │ store_id             │ INDEX        │ 中（门店订单管理）       │",
    "│ food_order       │ order_no             │ UNIQUE       │ 极高（唯一）           │",
    "│ food_order_item  │ order_id             │ INDEX        │ 高（订单详情 JOIN）     │",
    "│ payment_record   │ idempotency_key      │ UNIQUE       │ 极高（幂等防重）        │",
    "│ payment_record   │ order_id             │ INDEX        │ 高（支付查订单）        │",
    "│ cat              │ store_id             │ INDEX        │ 中（门店猫咪列表）       │",
    "│ cat_health_record│ cat_id               │ INDEX        │ 中（猫咪健康记录）       │",
    "│ user_preference  │ user_id              │ INDEX        │ 高（推荐引擎）          │",
    "│ review           │ order_id             │ UNIQUE       │ 极高（每单一评）        │",
    "│ refund_request   │ order_id             │ INDEX        │ 高（退款查订单）        │",
    "│ waiting_queue    │ store_id             │ INDEX        │ 高（门店队列）          │",
    "│ ticket           │                      │              │                      │",
    "│ points_transaction│ user_id             │ INDEX        │ 中（积分流水）          │",
    "│ activity_store   │ (activity_id,       │ UNIQUE       │ 极高（防重分发）        │",
    "│                  │  store_id)           │              │                      │",
    "│ reward_redemption│ (user_id,           │ UNIQUE       │ 极高（防重兑换）        │",
    "│                  │  activity_id)        │              │                      │",
    "└──────────────────┴──────────────────────┴──────────────┴──────────────────────┘",
    "",
    "5.2 性能预估",
    "数据规模预估（单门店运营 1 年）：",
    "• user: ~10,000 条（B+树索引，查询 < 1ms）",
    "• reservation: ~50,000 条（按月分区建议，当前暂不分表）",
    "• food_order: ~100,000 条（按 user_id 索引查询 < 5ms）",
    "• food_order_item: ~300,000 条（按 order_id 索引查询 < 3ms）",
    "• payment_record: ~100,000 条（idempotency_key 唯一索引查询 < 1ms）",
    "• points_transaction: ~200,000 条（按 user_id+created_at 复合索引查询 < 10ms）",
    "",
    "Top 5 慢查询治理思路：",
    "1. 看板聚合查询（GROUP BY store_id, DATE）：≈1s → 使用 dashboard_stat 预聚合表 + 定时任务刷新（已实现）",
    "2. 推荐评分查询（N门店 × M猫咪）：≈100ms → 限制门店范围（LBS 半径10km）+ 缓存用户偏好",
    "3. 顾客首页聚合（多表 JOIN）：≈50ms → 拆分并行查询 + 前端并发请求",
    "4. 退款审批列表（多状态过滤）：≈30ms → 复合索引 (store_id, status, created_at)",
    "5. 积分流水查询（排序+分页）：≈20ms → 复合索引 (user_id, created_at DESC)"
])

cs(doc.tables[14].cell(0,0),
    "以上索引与性能设计覆盖了 20+ 个关键索引、数据规模预估和 Top 5 慢查询治理思路，确保在 1000 并发预约场景下核心查询保持 P95 ≤ 350ms。")

# Section 6: 安全与合规 (T15, T16, T17, T18)
cm(doc.tables[15].cell(0,0), [
    "6.1 权限设计",
    "数据库账号与权限（最小权限原则）：",
    "",
    "┌──────────────────┬──────────────────────────────────────────────┐",
    "│ 账号              │ 权限                                          │",
    "├──────────────────┼──────────────────────────────────────────────┤",
    "│ nekocafe_app     │ SELECT, INSERT, UPDATE, DELETE (所有业务表)    │",
    "│                  │ EXECUTE (存储过程)                            │",
    "│                  │ 无 DDL 权限 (CREATE/ALTER/DROP)              │",
    "│                  │ 无 GRANT 权限                                │",
    "├──────────────────┼──────────────────────────────────────────────┤",
    "│ nekocafe_migration│ SELECT, INSERT, UPDATE, DELETE (所有表)      │",
    "│ (Flyway)         │ CREATE, ALTER, DROP (DDL 迁移)               │",
    "│                  │ CREATE INDEX, DROP INDEX                     │",
    "│                  │ 仅 Flyway 迁移时使用，应用运行时不使用此账号    │",
    "├──────────────────┼──────────────────────────────────────────────┤",
    "│ root             │ ALL PRIVILEGES                                │",
    "│                  │ 仅 DBA 使用，应用绝不使用 root 账号连接        │",
    "└──────────────────┴──────────────────────────────────────────────┘",
    "",
    "连接信息通过环境变量注入：MYSQL_HOST, MYSQL_PORT, MYSQL_DATABASE, MYSQL_USERNAME, MYSQL_PASSWORD",
    "应用配置文件不包含真实密码，仅使用 ${MYSQL_PASSWORD} 占位符引用环境变量。"
])

cm(doc.tables[16].cell(0,0), [
    "6.2 敏感字段处理",
    "",
    "敏感字段识别与处理策略：",
    "┌──────────────────┬──────────┬──────────────────────────────────┐",
    "│ 字段              │ 敏感级别  │ 处理方式                          │",
    "├──────────────────┼──────────┼──────────────────────────────────┤",
    "│ password_hash    │ 高        │ BCrypt 哈希存储（Salt+多轮），    │",
    "│                  │          │ 不可逆，不记录日志                 │",
    "│ phone            │ 中        │ 日志中脱敏（138****1234）          │",
    "│ email            │ 中        │ 日志中脱敏（u***@example.com）    │",
    "│ contact_name     │ 中        │ 日志中脱敏（张**）                 │",
    "│ contact_phone    │ 中        │ 日志中脱敏（138****1234）          │",
    "│ idempotency_key  │ 低        │ 不脱敏（排查支付问题需要完整值）    │",
    "│ JWT_SECRET       │ 极高       │ 仅通过环境变量注入，不进入数据库   │",
    "│ AI_API_KEY       │ 极高       │ 仅通过环境变量注入，不进入数据库   │",
    "└──────────────────┴──────────┴──────────────────────────────────┘",
    "",
    "脱敏实现：日志框架（Logback）配置 PatternLayout 脱敏规则，在日志输出前替换敏感信息。"
])

cm(doc.tables[17].cell(0,0), [
    "6.3 审计与备份",
    "",
    "审计策略：",
    "• 关键操作审计记录（通过日志 + 数据库表）：",
    "  - 用户创建/停用（admin 操作）→ 记录操作人、时间、目标用户",
    "  - 角色变更 → 记录变更前后角色",
    "  - 退款审批（同意/拒绝）→ 记录审批人、时间、原因",
    "  - 店长分配/撤换 → 记录操作人、时间、目标门店",
    "  - 活动发布/删除 → 记录操作人和时间",
    "",
    "备份策略：",
    "• 每日全量备份：mysqldump --single-transaction --routines --triggers nekocafe > backup_YYYYMMDD.sql",
    "• 答辩前备份：在 D14 答辩前进行一次全量备份，确保演示数据安全",
    "• 迁移脚本备份：Flyway 迁移脚本（V001-V025）通过 Git 版本控制",
    "• 恢复演练：答辩前进行一次备份恢复演练，验证备份文件可用",
    "",
    "生产环境增强建议（课程设计阶段不做硬性要求）：",
    "• MySQL 主从复制（1 Master + 1 Slave），读写分离",
    "• Binlog 增量备份（每 15 分钟）",
    "• 备份文件异地存储（OSS/S3）"
])

cs(doc.tables[18].cell(0,0),
    "以上安全与合规设计通过三账号权限分离、敏感字段分级脱敏、审计日志和备份策略，保障数据安全性和合规性。")

# Section 7: 数据生命周期 (T19)
cm(doc.tables[19].cell(0,0), [
    "7 数据生命周期",
    "",
    "各核心表的数据保留与清理策略：",
    "",
    "┌──────────────────┬──────────────┬──────────────────────────────────┐",
    "│ 表名              │ 保留期        │ 清理策略                          │",
    "├──────────────────┼──────────────┼──────────────────────────────────┤",
    "│ user             │ 永久          │ 逻辑删除（deleted=1），不物理删除   │",
    "│ reservation      │ 1 年          │ 1 年前 + 已过期 → 归档表/删除     │",
    "│ food_order       │ 2 年          │ 2 年前 + 已完成 → 归档表          │",
    "│ food_order_item  │ 2 年          │ 随 food_order 归档               │",
    "│ payment_record   │ 2 年          │ 随 food_order 归档               │",
    "│ reservation_slot │ 30 天         │ 过期时段 → 标记 INACTIVE，保留数据 │",
    "│ waiting_queue_*  │ 7 天          │ 已完成/取消的票据 → 定期清理       │",
    "│ points_transaction│ 1 年         │ 1 年前 → 归档表                  │",
    "│ review           │ 永久          │ 与订单关联，保留评价历史           │",
    "│ refund_request   │ 2 年          │ 随 food_order 归档               │",
    "│ cat_health_record│ 永久          │ 猫咪健康历史永久保留               │",
    "│ dashboard_stat   │ 1 年          │ 按天汇总，1 年前 → 按月聚合后删除  │",
    "│ promotion_activity│ 1 年         │ 结束后保留 1 年 → 归档            │",
    "└──────────────────┴──────────────┴──────────────────────────────────┘",
    "",
    "归档策略：",
    "• 课程设计阶段：不实现自动归档，手动 SQL 导出历史数据",
    "• 生产环境建议：定时任务（Spring @Scheduled）+ 归档表（_archive 后缀），保留归档数据 5 年",
    "",
    "销毁策略：",
    "• 逻辑删除优先：除队列数据外，绝大多数表使用逻辑删除（deleted=1）而非物理删除",
    "• 物理删除条件：归档超过保留期的数据 + 用户明确要求删除个人数据 + 测试/演示环境重建",
    "• 销毁审批：生产环境数据销毁需 DBA + 管理员双人确认"
])

doc.save(DST)
print(f"D-05 saved to {DST}")
